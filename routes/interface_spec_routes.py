"""
接口文档智能对照 - 路由层
所有 API 挂载在 /api 前缀下
"""
import os
import logging
from flask import Blueprint, request, current_app
from werkzeug.utils import secure_filename
from services.interface_parser_service import interface_parser
from services.interface_comparison_service import comparison_service
from services.interface_chat_service import interface_chat_service
from database import DatabasePool
from api_utils import api_response

logger = logging.getLogger(__name__)

spec_bp = Blueprint('interface_spec', __name__, url_prefix='/api')


# ========== 0. 文件文本提取（修复前端 404）==========

@spec_bp.route('/extract-text', methods=['POST'])
def extract_text_from_file():
    """
    从上传的 PDF/Word/TXT 文件中提取纯文本。
    前端 handleFileSelect() 调用此接口。
    """
    if 'file' not in request.files:
        return api_response(False, message='未找到上传文件')

    file = request.files['file']
    if not file or not file.filename:
        return api_response(False, message='文件为空')

    filename = secure_filename(file.filename)
    ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''

    # 保存临时文件
    upload_folder = current_app.config.get('UPLOAD_FOLDER', 'uploads')
    os.makedirs(upload_folder, exist_ok=True)
    tmp_path = os.path.join(upload_folder, f'_extract_tmp_{filename}')

    try:
        file.save(tmp_path)

        text = ''
        if ext == 'txt':
            # 尝试多种编码
            for encoding in ['utf-8', 'gbk', 'gb2312', 'latin-1']:
                try:
                    with open(tmp_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    break
                except UnicodeDecodeError:
                    continue

        elif ext == 'pdf':
            try:
                import pdfplumber
                with pdfplumber.open(tmp_path) as pdf:
                    pages = []
                    for page in pdf.pages:
                        page_text = page.extract_text() or ''
                        # 也提取表格
                        tables = page.extract_tables()
                        for table in tables:
                            for row in table:
                                if row:
                                    page_text += '\n' + '\t'.join([str(cell or '') for cell in row])
                        pages.append(page_text)
                    text = '\n\n'.join(pages)
            except ImportError:
                return api_response(False, message='服务器未安装 pdfplumber，请执行 pip install pdfplumber')
            except Exception as e:
                return api_response(False, message=f'PDF 解析失败: {str(e)}')

        elif ext in ('doc', 'docx'):
            try:
                from docx import Document
                doc = Document(tmp_path)
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                # 也提取表格
                for table in doc.tables:
                    for row in table.rows:
                        row_text = '\t'.join([cell.text.strip() for cell in row.cells])
                        if row_text.strip():
                            paragraphs.append(row_text)
                text = '\n'.join(paragraphs)
            except ImportError:
                return api_response(False, message='服务器未安装 python-docx，请执行 pip install python-docx')
            except Exception as e:
                return api_response(False, message=f'Word 文档解析失败: {str(e)}')

        elif ext in ('xml', 'wsdl'):
            with open(tmp_path, 'r', encoding='utf-8') as f:
                text = f.read()

        elif ext == 'json':
            with open(tmp_path, 'r', encoding='utf-8') as f:
                text = f.read()

        else:
            return api_response(False, message=f'不支持的文件格式: {ext}，支持 PDF/Word/TXT/XML/JSON')

        if not text.strip():
            return api_response(False, message='文件中未提取到文本内容，请检查文件是否为空或扫描件')

        return api_response(True, {
            'text': text,
            'length': len(text),
            'filename': filename
        })

    except Exception as e:
        logger.error(f"文件文本提取异常: {e}", exc_info=True)
        return api_response(False, message=f'文件处理失败: {str(e)}')
    finally:
        # 清理临时文件
        try:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)
        except:
            pass


# ========== 1. 文档解析 ==========

@spec_bp.route('/projects/<int:project_id>/interface-specs/parse', methods=['POST'])
def parse_interface_doc(project_id):
    """
    上传接口文档文本并 AI 解析为结构化接口。
    Body JSON: {
        "doc_text": "文档全文内容",
        "spec_source": "our_standard" 或 "vendor",
        "vendor_name": "厂商名(对方文档时填)",
        "doc_id": 可选, 关联 project_documents 表的 id
    }
    """
    data = request.json or {}
    doc_text = data.get('doc_text', '').strip()
    if not doc_text:
        return api_response(False, message='文档内容不能为空')

    spec_source = data.get('spec_source', 'vendor')
    category = data.get('category')
    vendor_name = data.get('vendor_name', '')
    doc_id = data.get('doc_id')

    save_project_id = None if spec_source == 'our_standard' and data.get('as_global') else project_id

    parsed = interface_parser.parse_document_with_ai(doc_text, spec_source, vendor_name)
    if not parsed:
        return api_response(False, message='AI 解析未能提取到接口定义，请检查文档内容或格式')

    created_ids = interface_parser.save_parsed_specs(
        save_project_id, doc_id, parsed, spec_source, vendor_name, category, raw_text=doc_text
    )

    return api_response(True, {
        'parsed_count': len(parsed),
        'spec_ids': created_ids,
        'interfaces': [{
            'name': p.get('interface_name', ''),
            'transcode': p.get('transcode', ''),
            'system_type': p.get('system_type', ''),
            'fields_count': len(p.get('fields', []))
        } for p in parsed]
    })


@spec_bp.route('/interface-specs/parse-standard', methods=['POST'])
def parse_our_standard():
    """上传并解析我方标准接口文档（全局标准，不绑定项目）"""
    data = request.json or {}
    doc_text = data.get('doc_text', '').strip()
    if not doc_text:
        return api_response(False, message='文档内容不能为空')

    parsed = interface_parser.parse_document_with_ai(doc_text, 'our_standard')
    if not parsed:
        return api_response(False, message='解析失败')

    created_ids = interface_parser.save_parsed_specs(
        None, data.get('doc_id'), parsed, 'our_standard',
        category=data.get('category'), raw_text=doc_text
    )
    return api_response(True, {
        'parsed_count': len(parsed),
        'spec_ids': created_ids,
        'interfaces': [{
            'name': p.get('interface_name', ''),
            'transcode': p.get('transcode', ''),
            'fields_count': len(p.get('fields', []))
        } for p in parsed]
    })


# ========== 2. 接口规范查询 ==========

@spec_bp.route('/projects/<int:project_id>/interface-specs', methods=['GET'])
def get_interface_specs(project_id):
    """获取项目的所有已解析接口规范(含字段)"""
    source = request.args.get('source')
    category = request.args.get('category')
    specs = interface_parser.get_specs_by_project(project_id, source, category)
    return api_response(True, specs)


@spec_bp.route('/interface-specs/standard', methods=['GET'])
def get_standard_specs():
    """获取全局标准接口规范"""
    category = request.args.get('category')
    specs = interface_parser.get_specs_by_project(None, 'our_standard', category)
    return api_response(True, specs)


@spec_bp.route('/interface-specs/<int:spec_id>', methods=['DELETE'])
def delete_spec(spec_id):
    """删除一个接口规范及其字段"""
    interface_parser.delete_spec(spec_id)
    return api_response(True, message='已删除')


# ========== 3. 智能对照 ==========

@spec_bp.route('/projects/<int:project_id>/interface-comparison/run', methods=['POST'])
def run_comparison(project_id):
    """一键执行接口对照（我方标准 vs 该项目的对方文档）"""
    data = request.json or {}
    category = data.get('category')
    result = comparison_service.run_full_comparison(project_id, category)
    return api_response(True, result)


@spec_bp.route('/projects/<int:project_id>/interface-comparisons', methods=['GET'])
def get_comparisons(project_id):
    """获取项目所有对照结果概览"""
    category = request.args.get('category')
    with DatabasePool.get_connection() as conn:
        query = '''
            SELECT ic.id, ic.our_spec_id, ic.vendor_spec_id, ic.match_type,
                   ic.match_confidence, ic.gap_count, ic.transform_count,
                   ic.status, ic.summary, ic.created_at, ic.category,
                   os.interface_name as our_name, os.transcode as our_transcode, os.system_type,
                   vs.interface_name as vendor_name, vs.transcode as vendor_transcode,
                   vs.vendor_name as vendor_company
            FROM interface_comparisons ic
            LEFT JOIN interface_specs os ON ic.our_spec_id = os.id
            LEFT JOIN interface_specs vs ON ic.vendor_spec_id = vs.id
            WHERE ic.project_id = ?
        '''
        params = [project_id]
        if category:
            query += ' AND ic.category = ?'
            params.append(category)

        query += ' ORDER BY ic.gap_count DESC, os.system_type'
        rows = conn.execute(query, params).fetchall()
    return api_response(True, [dict(r) for r in rows])


@spec_bp.route('/interface-comparisons/<int:comp_id>/detail', methods=['GET'])
def get_comparison_detail(comp_id):
    """获取单个对照的字段级映射详情"""
    with DatabasePool.get_connection() as conn:
        comp = conn.execute('SELECT * FROM interface_comparisons WHERE id = ?', (comp_id,)).fetchone()
        if not comp:
            return api_response(False, message='对照记录不存在', code=404)
        comp = dict(comp)
        mappings = [dict(m) for m in conn.execute(
            'SELECT * FROM field_mappings WHERE comparison_id = ? ORDER BY mapping_status, our_field_name',
            (comp_id,)
        ).fetchall()]
    comp['mappings'] = mappings
    return api_response(True, comp)


# ========== 4. AI 报告 ==========

@spec_bp.route('/projects/<int:project_id>/interface-comparison/report', methods=['GET'])
def get_comparison_report(project_id):
    """AI 生成该项目的接口对照分析报告"""
    report = comparison_service.generate_ai_report(project_id)
    return api_response(True, {'report': report})


# ========== 5. 字段映射确认 ==========

@spec_bp.route('/field-mappings/<int:mapping_id>/confirm', methods=['PUT'])
def confirm_mapping(mapping_id):
    """工程师手动确认/修正字段映射"""
    data = request.json or {}
    with DatabasePool.get_connection() as conn:
        updates = []
        params = []
        if 'mapping_status' in data:
            updates.append('mapping_status = ?')
            params.append(data['mapping_status'])
        if 'transform_rule' in data:
            updates.append('transform_rule = ?')
            params.append(data['transform_rule'])
        if 'remark' in data:
            updates.append('remark = ?')
            params.append(data['remark'])
        updates.append('is_confirmed = 1')
        params.append(mapping_id)

        conn.execute(f"UPDATE field_mappings SET {', '.join(updates)} WHERE id = ?", params)
        conn.commit()
    return api_response(True, message='已确认')


@spec_bp.route('/interface-comparisons/<int:comp_id>/status', methods=['PUT'])
def update_comparison_status(comp_id):
    """更新对照结果的审核状态"""
    data = request.json or {}
    with DatabasePool.get_connection() as conn:
        conn.execute('''
            UPDATE interface_comparisons SET status = ?, reviewed_by = ?,
            reviewed_at = CURRENT_TIMESTAMP WHERE id = ?
        ''', (data.get('status', 'reviewed'), data.get('reviewed_by', ''), comp_id))
        conn.commit()
    return api_response(True)


# ========== 6. AI 对话 + 请求生成（全新）==========

@spec_bp.route('/projects/<int:project_id>/interface-specs/chat', methods=['POST'])
def interface_chat(project_id):
    """
    接口 AI 助手对话。
    Body JSON: {
        "message": "用户提问",
        "category": "手麻标准" | "重症标准"
    }
    支持的意图：
    - 生成请求（XML/JSON/SQL）
    - 查询字段映射
    - 对接指导
    - 自由问答
    """
    data = request.json or {}
    message = data.get('message', '').strip()
    if not message:
        return api_response(False, message='消息不能为空')

    category = data.get('category', '手麻标准')

    try:
        result = interface_chat_service.chat(project_id, message, category)
        return api_response(True, result)
    except Exception as e:
        logger.error(f"接口 AI 对话异常: {e}", exc_info=True)
        return api_response(False, message=f'AI 助手响应失败: {str(e)}')


@spec_bp.route('/projects/<int:project_id>/interface-specs/generate-request', methods=['POST'])
def generate_request(project_id):
    """
    根据对照结果，为指定接口生成可复制的请求内容。
    Body JSON: {
        "comparison_id": 对照记录 ID,
        "format": "xml" | "json" | "sql" | "auto",
        "params": { 可选的请求参数 }
    }
    """
    data = request.json or {}
    comparison_id = data.get('comparison_id')
    if not comparison_id:
        return api_response(False, message='请指定对照记录 ID')

    req_format = data.get('format', 'auto')
    params = data.get('params', {})

    try:
        result = interface_chat_service.generate_request(
            project_id, comparison_id, req_format, params
        )
        return api_response(True, result)
    except Exception as e:
        logger.error(f"请求生成异常: {e}", exc_info=True)
        return api_response(False, message=f'请求生成失败: {str(e)}')
