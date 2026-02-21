# services/alignment_service.py
"""
智能接口文档对齐服务
职责：标准库管理、文档解析、AI对齐、请求生成
"""

import json
import os
import re
import xml.etree.ElementTree as ET
from datetime import datetime
from database import DatabasePool
from services.ai_service import ai_service


class AlignmentService:
    """接口文档对齐核心服务"""

    # ================================================================
    #  第一层：标准接口库管理
    # ================================================================

    @staticmethod
    def get_spec_versions():
        """获取所有标准版本"""
        with DatabasePool.get_connection() as conn:
            rows = conn.execute('''
                SELECT spec_version, category,
                       COUNT(*) as interface_count,
                       MAX(created_at) as last_updated
                FROM interface_specs
                GROUP BY spec_version, category
                ORDER BY category, spec_version DESC
            ''').fetchall()
            return [dict(r) for r in rows]

    @staticmethod
    def get_spec_interfaces(spec_version):
        """获取某版本的全部标准接口及其字段"""
        with DatabasePool.get_connection() as conn:
            interfaces = conn.execute('''
                SELECT * FROM interface_specs
                WHERE spec_version = ?
                ORDER BY sort_order, id
            ''', (spec_version,)).fetchall()

            result = []
            for iface in interfaces:
                iface_dict = dict(iface)
                fields = conn.execute('''
                    SELECT * FROM interface_spec_fields
                    WHERE spec_interface_id = ?
                    ORDER BY sort_order, id
                ''', (iface_dict['id'],)).fetchall()
                iface_dict['fields'] = [dict(f) for f in fields]
                result.append(iface_dict)
            return result

    @staticmethod
    def save_spec_interface(data):
        """保存/更新一个标准接口定义（含字段）"""
        with DatabasePool.get_connection() as conn:
            cursor = conn.cursor()

            interface_id = data.get('id')
            if interface_id:
                # 更新
                cursor.execute('''
                    UPDATE interface_specs SET
                        spec_version=?, category=?, system_name=?,
                        interface_name=?, interface_code=?, description=?,
                        protocol=?, view_name=?, is_required=?, sort_order=?
                    WHERE id=?
                ''', (
                    data['spec_version'], data['category'], data['system_name'],
                    data['interface_name'], data.get('interface_code', ''),
                    data.get('description', ''), data.get('protocol', '视图'),
                    data.get('view_name', ''), data.get('is_required', 0),
                    data.get('sort_order', 0), interface_id
                ))
            else:
                # 新建
                cursor.execute('''
                    INSERT INTO interface_specs
                    (spec_version, category, system_name, interface_name,
                     interface_code, description, protocol, view_name,
                     is_required, sort_order)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                ''', (
                    data['spec_version'], data['category'], data['system_name'],
                    data['interface_name'], data.get('interface_code', ''),
                    data.get('description', ''), data.get('protocol', '视图'),
                    data.get('view_name', ''), data.get('is_required', 0),
                    data.get('sort_order', 0)
                ))
                interface_id = cursor.lastrowid

            # 保存字段 —— 先删后插
            cursor.execute(
                'DELETE FROM interface_spec_fields WHERE spec_interface_id=?',
                (interface_id,)
            )
            for idx, field in enumerate(data.get('fields', [])):
                cursor.execute('''
                    INSERT INTO interface_spec_fields
                    (spec_interface_id, field_name, field_label, field_type,
                     is_required, max_length, sample_value, remark, sort_order)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                ''', (
                    interface_id,
                    field['field_name'], field.get('field_label', ''),
                    field.get('field_type', 'VARCHAR'),
                    field.get('is_required', 0), field.get('max_length'),
                    field.get('sample_value', ''), field.get('remark', ''), idx
                ))

            conn.commit()
            return interface_id

    @staticmethod
    def delete_spec_interface(spec_id):
        with DatabasePool.get_connection() as conn:
            conn.execute('DELETE FROM interface_spec_fields WHERE spec_interface_id=?', (spec_id,))
            conn.execute('DELETE FROM interface_specs WHERE id=?', (spec_id,))
            conn.commit()

    @staticmethod
    def import_spec_from_file(file_path, spec_version, category):
        """从上传的文档中 AI 解析出标准接口并入库"""
        raw_text = AlignmentService._extract_text(file_path)
        if not raw_text:
            return {'success': False, 'message': '文档解析失败，内容为空'}

        interfaces = AlignmentService._ai_extract_spec(raw_text, spec_version, category)
        if not interfaces:
            return {'success': False, 'message': 'AI 未能从文档中识别到接口'}

        count = 0
        for iface in interfaces:
            iface['spec_version'] = spec_version
            iface['category'] = category
            AlignmentService.save_spec_interface(iface)
            count += 1

        return {'success': True, 'count': count}

    # ================================================================
    #  第二层：对方文档解析
    # ================================================================

    @staticmethod
    def _extract_text(file_path):
        """从文件中提取文本内容，支持多种格式"""
        if not os.path.exists(file_path):
            return None
        ext = os.path.splitext(file_path)[1].lower()

        try:
            if ext == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()

            elif ext == '.pdf':
                try:
                    import pdfplumber
                    parts = []
                    with pdfplumber.open(file_path) as pdf:
                        for page in pdf.pages:
                            tables = page.extract_tables()
                            if tables:
                                for table in tables:
                                    for row in table:
                                        parts.append(' | '.join(
                                            [str(c or '') for c in row]))
                            else:
                                t = page.extract_text()
                                if t:
                                    parts.append(t)
                    return '\n'.join(parts)
                except ImportError:
                    return None

            elif ext in ('.docx', '.doc'):
                from docx import Document
                doc = Document(file_path)
                parts = []
                for p in doc.paragraphs:
                    if p.text.strip():
                        parts.append(p.text)
                for table in doc.tables:
                    for row in table.rows:
                        cells = [c.text.strip() for c in row.cells]
                        parts.append(' | '.join(cells))
                return '\n'.join(parts)

            elif ext == '.xml':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()

            elif ext == '.json':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()

            elif ext == '.wsdl':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()

            else:
                # 尝试当文本读
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
        except Exception as e:
            print(f"文件解析失败 [{file_path}]: {e}")
            return None

    @staticmethod
    def _detect_format(text):
        """自动检测文档内容的技术格式"""
        text_stripped = text.strip()
        if text_stripped.startswith('{') or text_stripped.startswith('['):
            try:
                json.loads(text_stripped)
                return 'json'
            except:
                pass
        if text_stripped.startswith('<?xml') or text_stripped.startswith('<'):
            if 'wsdl:' in text or 'definitions' in text[:500]:
                return 'wsdl'
            return 'xml'
        if 'MSH|' in text[:200] or 'PID|' in text[:500]:
            return 'hl7'
        if 'CREATE VIEW' in text.upper() or 'SELECT' in text.upper()[:200]:
            return 'sql'
        if 'swagger' in text[:500].lower() or 'openapi' in text[:500].lower():
            return 'swagger'
        return 'text'

    @staticmethod
    def _programmatic_parse_json(text):
        """程序化解析 JSON/Swagger 格式接口"""
        try:
            data = json.loads(text)
        except:
            return None

        interfaces = []

        # Swagger / OpenAPI 格式
        if 'paths' in data:
            for path, methods in data.get('paths', {}).items():
                for method, detail in methods.items():
                    if method in ('get', 'post', 'put', 'delete'):
                        fields = []
                        for param in detail.get('parameters', []):
                            fields.append({
                                'field_name': param.get('name', ''),
                                'field_label': param.get('description', ''),
                                'field_type': param.get('type',
                                              param.get('schema', {}).get('type', 'string')),
                                'is_required': param.get('required', False),
                            })
                        interfaces.append({
                            'interface_name': detail.get('summary',
                                              detail.get('operationId', path)),
                            'view_name': f"{method.upper()} {path}",
                            'system_name': data.get('info', {}).get('title', '未知系统'),
                            'protocol': 'REST',
                            'description': detail.get('description', ''),
                            'fields': fields,
                        })
            return interfaces if interfaces else None

        # 普通 JSON 数组（接口列表）
        if isinstance(data, list):
            for item in data:
                if isinstance(item, dict) and ('interface_name' in item or
                                                'name' in item or
                                                'method' in item):
                    fields = []
                    for f in item.get('fields', item.get('parameters', [])):
                        if isinstance(f, dict):
                            fields.append({
                                'field_name': f.get('field_name',
                                              f.get('name', '')),
                                'field_label': f.get('field_label',
                                              f.get('label',
                                              f.get('description', ''))),
                                'field_type': f.get('field_type',
                                              f.get('type', 'VARCHAR')),
                                'is_required': f.get('is_required',
                                              f.get('required', False)),
                            })
                    interfaces.append({
                        'interface_name': item.get('interface_name',
                                          item.get('name', '')),
                        'view_name': item.get('view_name',
                                     item.get('method', '')),
                        'system_name': item.get('system_name', '未知系统'),
                        'protocol': item.get('protocol', 'REST'),
                        'description': item.get('description',
                                       item.get('remark', '')),
                        'fields': fields,
                    })
            return interfaces if interfaces else None

        return None

    @staticmethod
    def _programmatic_parse_xml(text):
        """程序化解析 XML/WSDL"""
        interfaces = []
        try:
            # 去除 BOM
            if text.startswith('\ufeff'):
                text = text[1:]
            root = ET.fromstring(text)
        except ET.ParseError:
            return None

        ns = {
            'wsdl': 'http://schemas.xmlsoap.org/wsdl/',
            'xs': 'http://www.w3.org/2001/XMLSchema',
            'xsd': 'http://www.w3.org/2001/XMLSchema',
            'soap': 'http://schemas.xmlsoap.org/wsdl/soap/',
        }

        # WSDL 格式
        operations = root.findall('.//wsdl:operation', ns)
        if not operations:
            # 尝试无命名空间
            operations = root.findall('.//{http://schemas.xmlsoap.org/wsdl/}operation')

        if operations:
            for op in operations:
                name = op.get('name', '未命名')
                doc_el = op.find('.//wsdl:documentation', ns)
                if doc_el is None:
                    doc_el = op.find('.//{http://schemas.xmlsoap.org/wsdl/}documentation')
                description = doc_el.text if doc_el is not None else ''
                interfaces.append({
                    'interface_name': name,
                    'view_name': name,
                    'system_name': '未知系统',
                    'protocol': 'WebService',
                    'description': description,
                    'fields': [],
                })
            # 解析 types 中的 element 获取字段
            elements = root.findall('.//xs:element', ns)
            if not elements:
                elements = root.findall(
                    './/{http://www.w3.org/2001/XMLSchema}element')
            el_map = {}
            for el in elements:
                el_name = el.get('name', '')
                el_type = el.get('type', '')
                parent = el
                # 往上找到属于哪个 complexType
                # 简化处理：收集所有 element
                el_map[el_name] = el_type

            return interfaces if interfaces else None

        # 普通 XML —— 尝试按表格/条目结构提取
        # 通用策略：找到所有包含子元素的节点
        for child in root:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            fields = []
            for sub in child:
                sub_tag = sub.tag.split('}')[-1] if '}' in sub.tag else sub.tag
                fields.append({
                    'field_name': sub_tag,
                    'field_label': '',
                    'field_type': 'VARCHAR',
                    'is_required': False,
                    'remark': (sub.text or '')[:100],
                })
            if fields:
                interfaces.append({
                    'interface_name': tag,
                    'view_name': '',
                    'system_name': '未知系统',
                    'protocol': 'XML',
                    'description': '',
                    'fields': fields,
                })

        return interfaces if interfaces else None

    @staticmethod
    def _programmatic_parse_hl7(text):
        """解析 HL7 v2 消息"""
        interfaces = []
        segments = {}
        for line in text.split('\n'):
            line = line.strip()
            if '|' in line:
                seg_id = line.split('|')[0]
                if seg_id and len(seg_id) == 3 and seg_id.isalpha():
                    if seg_id not in segments:
                        segments[seg_id] = []
                    parts = line.split('|')
                    segments[seg_id].append(parts)

        HL7_SEGMENT_NAMES = {
            'MSH': '消息头', 'PID': '患者信息', 'PV1': '就诊信息',
            'OBR': '检验/检查请求', 'OBX': '检验/检查结果',
            'ORC': '医嘱控制', 'RXA': '用药管理', 'DG1': '诊断信息',
            'IN1': '保险信息', 'NK1': '联系人', 'AL1': '过敏信息',
            'SCH': '排程信息',
        }

        for seg_id, rows in segments.items():
            sample_row = rows[0] if rows else []
            fields = []
            for i, val in enumerate(sample_row):
                if i == 0:
                    continue  # 段名本身
                fields.append({
                    'field_name': f'{seg_id}-{i}',
                    'field_label': '',
                    'field_type': 'VARCHAR',
                    'is_required': False,
                    'remark': (val or '')[:80],
                })
            interfaces.append({
                'interface_name': HL7_SEGMENT_NAMES.get(seg_id,
                                   f'HL7段-{seg_id}'),
                'view_name': seg_id,
                'system_name': 'HL7消息',
                'protocol': 'HL7v2',
                'description': f'HL7 {seg_id} 段，共 {len(fields)} 个字段',
                'fields': fields,
            })

        return interfaces if interfaces else None

    @staticmethod
    def _programmatic_parse_sql(text):
        """解析 SQL 视图/表定义"""
        interfaces = []
        # 匹配 CREATE VIEW xxx AS 或 CREATE TABLE xxx
        pattern = re.compile(
            r'CREATE\s+(?:OR\s+REPLACE\s+)?(?:VIEW|TABLE)\s+'
            r'[`"\[]?(\w+)[`"\]]?\s*(?:\(|AS)',
            re.IGNORECASE
        )
        for match in pattern.finditer(text):
            view_name = match.group(1)
            # 尝试提取字段
            start = match.end()
            # 找到对应的括号内容或 SELECT 列
            rest = text[start:start + 2000]
            fields = []
            # 简单匹配列名
            col_pattern = re.compile(
                r'[`"\[]?(\w+)[`"\]]?\s+(VARCHAR|INT|INTEGER|DATE|DATETIME|'
                r'TEXT|FLOAT|REAL|NUMERIC|TIMESTAMP|CHAR|NVARCHAR|NUMBER)',
                re.IGNORECASE
            )
            for col in col_pattern.finditer(rest):
                fields.append({
                    'field_name': col.group(1),
                    'field_label': '',
                    'field_type': col.group(2).upper(),
                    'is_required': False,
                })

            interfaces.append({
                'interface_name': view_name,
                'view_name': view_name,
                'system_name': '数据库视图',
                'protocol': '视图/SQL',
                'description': '',
                'fields': fields,
            })

        return interfaces if interfaces else None

    @staticmethod
    def parse_vendor_document(file_path=None, raw_text=None):
        """
        解析第三方接口文档 —— 自动检测格式并选择策略
        返回结构化的接口列表
        """
        if file_path:
            text = AlignmentService._extract_text(file_path)
        elif raw_text:
            text = raw_text
        else:
            return []

        if not text or not text.strip():
            return []

        fmt = AlignmentService._detect_format(text)

        # 优先尝试程序化解析（准确性高、速度快、不消耗 AI 额度）
        result = None
        if fmt == 'json' or fmt == 'swagger':
            result = AlignmentService._programmatic_parse_json(text)
        elif fmt in ('xml', 'wsdl'):
            result = AlignmentService._programmatic_parse_xml(text)
        elif fmt == 'hl7':
            result = AlignmentService._programmatic_parse_hl7(text)
        elif fmt == 'sql':
            result = AlignmentService._programmatic_parse_sql(text)

        if result:
            return result

        # 程序化解析失败或格式为纯文本 → 走 AI 提取
        return AlignmentService._ai_extract_vendor(text)

    # ================================================================
    #  AI 调用层
    # ================================================================

    @staticmethod
    def _ai_extract_spec(text, spec_version, category):
        """AI 从我方文档中提取标准接口定义"""
        system_prompt = """你是一位资深医疗信息化工程师，精通ICU重症和手术麻醉信息系统。
请从用户提供的接口规范文档中，提取所有接口的结构化定义。

严格返回 JSON 数组：
[
  {
    "system_name": "HIS/LIS/PACS/EMR",
    "interface_name": "接口中文名称",
    "interface_code": "英文编码(如 icu_order)",
    "description": "用途说明",
    "protocol": "视图/WebService/REST/存储过程",
    "view_name": "视图名或方法名",
    "is_required": true/false,
    "sort_order": 序号,
    "fields": [
      {
        "field_name": "字段名",
        "field_label": "中文含义",
        "field_type": "VARCHAR/INT/DATE/DATETIME",
        "is_required": true/false,
        "max_length": 50,
        "sample_value": "示例值",
        "remark": "备注"
      }
    ]
  }
]

只返回 JSON，不要任何解释文字。"""

        user_content = text[:15000]  # 限制长度
        resp = ai_service.call_ai_api(system_prompt, user_content, task_type="json")
        return AlignmentService._safe_parse_json_array(resp)

    @staticmethod
    def _ai_extract_vendor(text):
        """AI 从第三方非结构化文档中提取接口列表"""
        system_prompt = """你是一位资深医疗信息化集成工程师，精通 HIS/LIS/PACS/EMR 系统接口。
请从用户提供的第三方接口文档中，提取所有接口/视图的结构化信息。

严格返回 JSON 数组：
[
  {
    "interface_name": "接口名称",
    "view_name": "视图名/方法名/API路径",
    "system_name": "所属系统(HIS/LIS/PACS/EMR/其他)",
    "protocol": "协议(视图/WebService/REST/HL7/存储过程)",
    "description": "用途简述",
    "fields": [
      {
        "field_name": "字段名",
        "field_label": "中文含义",
        "field_type": "数据类型",
        "is_required": false,
        "remark": "备注"
      }
    ]
  }
]
如果文档中没有字段明细，fields 返回空数组。
只返回 JSON，不要任何解释文字。"""

        user_content = text[:15000]
        resp = ai_service.call_ai_api(system_prompt, user_content, task_type="json")
        return AlignmentService._safe_parse_json_array(resp)

    # ================================================================
    #  第三层：AI 对齐引擎
    # ================================================================

    @staticmethod
    def run_alignment(project_id, spec_version, vendor_name,
                      file_path=None, raw_text=None, created_by='system'):
        """
        完整对齐流程：解析 → 对齐 → 持久化
        """
        with DatabasePool.get_connection() as conn:
            cursor = conn.cursor()
            # 创建会话
            cursor.execute('''
                INSERT INTO alignment_sessions
                (project_id, spec_version, vendor_name,
                 vendor_doc_path, status, created_by)
                VALUES (?, ?, ?, ?, 'parsing', ?)
            ''', (project_id, spec_version, vendor_name,
                  file_path or '', created_by))
            session_id = cursor.lastrowid
            conn.commit()

        try:
            # Step 1: 解析对方文档
            vendor_interfaces = AlignmentService.parse_vendor_document(
                file_path=file_path, raw_text=raw_text
            )
            if not vendor_interfaces:
                AlignmentService._update_session_status(
                    session_id, 'failed', error='对方文档解析结果为空')
                return {'success': False, 'session_id': session_id,
                        'message': '未能从文档中解析到任何接口'}

            # 保存解析出的原始文本
            doc_text = ''
            if file_path:
                doc_text = AlignmentService._extract_text(file_path) or ''
            elif raw_text:
                doc_text = raw_text
            with DatabasePool.get_connection() as conn:
                conn.execute(
                    'UPDATE alignment_sessions SET vendor_doc_text=?, status=? WHERE id=?',
                    (doc_text[:50000], 'aligning', session_id))

            # Step 2: 加载我方标准
            spec_interfaces = AlignmentService.get_spec_interfaces(spec_version)
            if not spec_interfaces:
                AlignmentService._update_session_status(
                    session_id, 'failed',
                    error=f'标准版本 [{spec_version}] 中无接口定义')
                return {'success': False, 'session_id': session_id,
                        'message': f'标准 [{spec_version}] 中没有接口，请先维护标准库'}

            # Step 3: AI 对齐
            alignment_result = AlignmentService._ai_align(
                spec_interfaces, vendor_interfaces
            )

            # Step 4: 持久化
            AlignmentService._save_alignment_results(
                session_id, spec_interfaces, alignment_result
            )

            # Step 5: 更新会话统计
            AlignmentService._finalize_session(session_id, alignment_result)

            return {
                'success': True,
                'session_id': session_id,
                'summary': alignment_result.get('summary', ''),
                'risk_assessment': alignment_result.get('risk_assessment', ''),
            }

        except Exception as e:
            print(f"对齐流程异常: {e}")
            import traceback
            traceback.print_exc()
            AlignmentService._update_session_status(
                session_id, 'failed', error=str(e))
            return {'success': False, 'session_id': session_id,
                    'message': f'对齐过程出错: {str(e)}'}

    @staticmethod
    def _ai_align(spec_interfaces, vendor_interfaces):
        """调用 AI 进行标准 vs 对方的对齐比对"""
        # 精简数据发给 AI，避免超 token
        spec_summary = []
        for s in spec_interfaces:
            fields_brief = [f.get('field_name', '') for f in s.get('fields', [])[:20]]
            spec_summary.append({
                'id': s['id'],
                'code': s.get('interface_code', ''),
                'name': s['interface_name'],
                'system': s['system_name'],
                'protocol': s.get('protocol', ''),
                'view': s.get('view_name', ''),
                'required': bool(s.get('is_required')),
                'fields': fields_brief,
            })

        vendor_summary = []
        for i, v in enumerate(vendor_interfaces[:50]):
            fields_brief = [f.get('field_name', '') for f in v.get('fields', [])[:20]]
            vendor_summary.append({
                'idx': i,
                'name': v.get('interface_name', ''),
                'view': v.get('view_name', ''),
                'system': v.get('system_name', ''),
                'protocol': v.get('protocol', ''),
                'desc': (v.get('description', '') or '')[:100],
                'fields': fields_brief,
            })

        system_prompt = """你是医疗信息化接口集成专家。将"我方标准接口"与"第三方接口"逐条匹配。

匹配规则：
- matched: 用途高度一致，可直接对接
- partial: 用途相似但名称/字段有差异
- missing: 我方需要但对方未提供

返回 JSON（不要任何多余文字）：
{
  "alignments": [
    {
      "spec_id": 我方接口ID(数字),
      "match_status": "matched/partial/missing",
      "confidence": 0.0到1.0,
      "vendor_idx": 对方接口序号(missing时为null),
      "vendor_interface_name": "对方接口名(missing时为null)",
      "vendor_view_name": "对方视图名",
      "diff_summary": "差异说明",
      "risk_note": "风险提示(必选接口缺失时务必标注)",
      "field_mappings": [
        {"spec_field": "我方字段名", "vendor_field": "对方字段名", "transform": "转换说明或null"}
      ]
    }
  ],
  "extras": [
    {"vendor_idx": 对方序号, "name": "对方多出的接口名", "suggestion": "建议"}
  ],
  "summary": "一段话总结对齐情况(中文)",
  "risk_assessment": "关键风险点(中文)"
}"""

        user_content = f"""## 我方标准接口 ({len(spec_summary)}个)
{json.dumps(spec_summary, ensure_ascii=False)}

## 第三方接口 ({len(vendor_summary)}个)
{json.dumps(vendor_summary, ensure_ascii=False)}"""

        resp = ai_service.call_ai_api(system_prompt, user_content, task_type="json")
        result = AlignmentService._safe_parse_json(resp)
        if not result:
            result = {'alignments': [], 'extras': [], 'summary': 'AI 解析失败',
                      'risk_assessment': ''}
        return result

    @staticmethod
    def _save_alignment_results(session_id, spec_interfaces, ai_result):
        """将 AI 对齐结果存入数据库"""
        spec_map = {s['id']: s for s in spec_interfaces}

        with DatabasePool.get_connection() as conn:
            cursor = conn.cursor()

            for item in ai_result.get('alignments', []):
                spec_id = item.get('spec_id')
                if spec_id not in spec_map:
                    continue

                cursor.execute('''
                    INSERT INTO alignment_results
                    (session_id, spec_interface_id, match_status, confidence,
                     vendor_interface_name, vendor_view_name, vendor_protocol,
                     vendor_description, diff_summary, risk_note)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                ''', (
                    session_id, spec_id,
                    item.get('match_status', 'missing'),
                    item.get('confidence', 0),
                    item.get('vendor_interface_name'),
                    item.get('vendor_view_name'),
                    item.get('vendor_protocol'),
                    item.get('vendor_description'),
                    item.get('diff_summary', ''),
                    item.get('risk_note', ''),
                ))
                result_id = cursor.lastrowid

                # 保存字段映射
                for fm in item.get('field_mappings', []):
                    # 查找 spec_field_id
                    spec_iface = spec_map.get(spec_id, {})
                    spec_field_id = None
                    for f in spec_iface.get('fields', []):
                        if f['field_name'] == fm.get('spec_field'):
                            spec_field_id = f['id']
                            break
                    if not spec_field_id:
                        continue

                    cursor.execute('''
                        INSERT INTO alignment_field_maps
                        (alignment_result_id, spec_field_id,
                         vendor_field_name, vendor_field_type,
                         map_status, transform_rule, confidence)
                        VALUES (?, ?, ?, ?, 'auto', ?, ?)
                    ''', (
                        result_id, spec_field_id,
                        fm.get('vendor_field', ''),
                        fm.get('vendor_type', ''),
                        fm.get('transform'),
                        item.get('confidence', 0),
                    ))

            # 处理对方多出的接口
            for extra in ai_result.get('extras', []):
                cursor.execute('''
                    INSERT INTO alignment_results
                    (session_id, spec_interface_id, match_status, confidence,
                     vendor_interface_name, diff_summary, risk_note)
                    VALUES (?, 0, 'extra', 0, ?, ?, ?)
                ''', (
                    session_id,
                    extra.get('name', ''),
                    extra.get('suggestion', ''),
                    '',
                ))

            conn.commit()

    @staticmethod
    def _finalize_session(session_id, ai_result):
        """统计并更新会话"""
        alignments = ai_result.get('alignments', [])
        matched = sum(1 for a in alignments if a.get('match_status') == 'matched')
        partial = sum(1 for a in alignments if a.get('match_status') == 'partial')
        missing = sum(1 for a in alignments if a.get('match_status') == 'missing')
        extras = len(ai_result.get('extras', []))
        total = len(alignments)
        score = round((matched * 100 + partial * 60) / total, 1) if total > 0 else 0

        with DatabasePool.get_connection() as conn:
            conn.execute('''
                UPDATE alignment_sessions SET
                    status='completed',
                    total_spec_interfaces=?,
                    matched_count=?, partial_count=?,
                    missing_count=?, extra_count=?,
                    match_score=?,
                    ai_summary=?,
                    completed_at=?
                WHERE id=?
            ''', (
                total, matched, partial, missing, extras, score,
                ai_result.get('summary', ''),
                datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                session_id
            ))

    @staticmethod
    def _update_session_status(session_id, status, error=''):
        with DatabasePool.get_connection() as conn:
            conn.execute(
                'UPDATE alignment_sessions SET status=?, ai_summary=? WHERE id=?',
                (status, error, session_id))

    # ================================================================
    #  查询对齐结果
    # ================================================================

    @staticmethod
    def get_sessions(project_id):
        """获取项目的所有对齐会话"""
        with DatabasePool.get_connection() as conn:
            rows = conn.execute('''
                SELECT * FROM alignment_sessions
                WHERE project_id = ?
                ORDER BY created_at DESC
            ''', (project_id,)).fetchall()
            return [dict(r) for r in rows]

    @staticmethod
    def get_session_detail(session_id):
        """获取对齐会话详情（含每条对齐结果）"""
        with DatabasePool.get_connection() as conn:
            session = conn.execute(
                'SELECT * FROM alignment_sessions WHERE id=?',
                (session_id,)
            ).fetchone()
            if not session:
                return None

            session_dict = dict(session)

            results = conn.execute('''
                SELECT ar.*, isp.interface_name as spec_interface_name,
                       isp.system_name as spec_system_name,
                       isp.interface_code as spec_code,
                       isp.is_required as spec_is_required,
                       isp.protocol as spec_protocol,
                       isp.view_name as spec_view_name
                FROM alignment_results ar
                LEFT JOIN interface_specs isp ON ar.spec_interface_id = isp.id
                WHERE ar.session_id = ?
                ORDER BY
                    CASE ar.match_status
                        WHEN 'missing' THEN 0
                        WHEN 'partial' THEN 1
                        WHEN 'matched' THEN 2
                        WHEN 'extra' THEN 3
                    END,
                    ar.confidence DESC
            ''', (session_id,)).fetchall()

            result_list = []
            for r in results:
                r_dict = dict(r)
                # 加载字段映射
                field_maps = conn.execute('''
                    SELECT afm.*, isf.field_name as spec_field_name,
                           isf.field_label as spec_field_label,
                           isf.field_type as spec_field_type,
                           isf.is_required as spec_field_required
                    FROM alignment_field_maps afm
                    LEFT JOIN interface_spec_fields isf ON afm.spec_field_id = isf.id
                    WHERE afm.alignment_result_id = ?
                    ORDER BY isf.sort_order
                ''', (r_dict['id'],)).fetchall()
                r_dict['field_maps'] = [dict(fm) for fm in field_maps]
                result_list.append(r_dict)

            session_dict['results'] = result_list
            return session_dict

    @staticmethod
    def confirm_result(result_id, confirmed_by, manual_note=''):
        """人工确认某条对齐结果"""
        with DatabasePool.get_connection() as conn:
            conn.execute('''
                UPDATE alignment_results SET
                    is_confirmed=1, confirmed_by=?, confirmed_at=?, manual_note=?
                WHERE id=?
            ''', (confirmed_by, datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                  manual_note, result_id))

    @staticmethod
    def update_field_map(map_id, vendor_field_name, transform_rule=''):
        """手动修正字段映射"""
        with DatabasePool.get_connection() as conn:
            conn.execute('''
                UPDATE alignment_field_maps SET
                    vendor_field_name=?, transform_rule=?, map_status='manual'
                WHERE id=?
            ''', (vendor_field_name, transform_rule, map_id))

    @staticmethod
    def delete_session(session_id):
        """删除对齐会话及其全部结果"""
        with DatabasePool.get_connection() as conn:
            # 先删字段映射
            conn.execute('''
                DELETE FROM alignment_field_maps
                WHERE alignment_result_id IN
                    (SELECT id FROM alignment_results WHERE session_id=?)
            ''', (session_id,))
            conn.execute('DELETE FROM alignment_results WHERE session_id=?',
                         (session_id,))
            conn.execute('DELETE FROM alignment_sessions WHERE id=?',
                         (session_id,))
            conn.commit()

    # ================================================================
    #  第四层：AI 请求生成器
    # ================================================================

    @staticmethod
    def ai_generate_request(session_id, question):
        """
        AI 对接助手：根据对齐结果回答工程师的问题
        支持：生成请求体、查字段映射、排错建议等
        """
        # 1. 加载上下文
        detail = AlignmentService.get_session_detail(session_id)
        if not detail:
            return {'answer': '对齐会话不存在', 'code_blocks': []}

        # 构建上下文摘要
        context_parts = [
            f"项目对接方: {detail['vendor_name']}",
            f"使用标准: {detail['spec_version']}",
            f"匹配度: {detail['match_score']}%",
            f"",
            "## 已对齐的接口映射关系：",
        ]

        for r in detail.get('results', []):
            if r['match_status'] == 'extra':
                continue
            line = (f"- 我方【{r.get('spec_interface_name', '?')}】"
                    f"({r.get('spec_protocol', '')}, "
                    f"视图: {r.get('spec_view_name', '')}) "
                    f"↔ 对方【{r.get('vendor_interface_name', '缺失')}】"
                    f"(视图: {r.get('vendor_view_name', '')})"
                    f" [{r['match_status']}]")
            context_parts.append(line)

            # 字段映射
            if r.get('field_maps'):
                for fm in r['field_maps']:
                    context_parts.append(
                        f"    {fm.get('spec_field_name', '?')}({fm.get('spec_field_label', '')}) "
                        f"↔ {fm.get('vendor_field_name', '未映射')}"
                        f"{' [需转换: ' + fm['transform_rule'] + ']' if fm.get('transform_rule') else ''}"
                    )

        context = '\n'.join(context_parts)

        system_prompt = f"""你是一位资深医疗信息化实施工程师，正在帮助同事进行接口对接工作。

以下是当前项目的接口对齐映射信息（已由 AI 自动比对生成）：
---
{context[:8000]}
---

请根据以上映射信息回答用户的问题。
规则：
1. 如果用户要求"生成请求"，根据对方接口的协议生成可直接使用的请求内容
   - WebService → 生成完整的 SOAP XML
   - REST → 生成 curl 命令或 HTTP 请求
   - 视图/SQL → 生成 SELECT 语句
   - HL7 → 生成 HL7 消息
2. 如果用户问"某字段对应什么"，从映射中查找并回答
3. 如果用户问排错问题，根据映射差异分析可能原因
4. 代码/请求内容用 ```xml ``` 或 ```sql ``` 等代码块包裹
5. 回答要简洁实用，直接可以复制使用
"""

        resp = ai_service.call_ai_api(system_prompt, question, task_type="code")
        if not resp:
            return {'answer': 'AI 暂时无法回答，请稍后重试', 'code_blocks': []}

        # 提取代码块方便前端单独显示复制按钮
        code_blocks = re.findall(r'```(\w*)\n(.*?)```', resp, re.DOTALL)
        blocks = [{'lang': cb[0] or 'text', 'code': cb[1].strip()}
                  for cb in code_blocks]

        return {'answer': resp, 'code_blocks': blocks}

    # ================================================================
    #  工具方法
    # ================================================================

    @staticmethod
    def _safe_parse_json_array(text):
        """安全解析 AI 返回的 JSON 数组"""
        if not text:
            return []
        try:
            text = text.strip()
            if '```' in text:
                text = text.split('```json')[-1].split('```')[0]
                text = text.strip()
            if not text.startswith('['):
                # 尝试找到第一个 [
                idx = text.find('[')
                if idx >= 0:
                    text = text[idx:]
            result = json.loads(text)
            return result if isinstance(result, list) else []
        except:
            return []

    @staticmethod
    def _safe_parse_json(text):
        """安全解析 AI 返回的 JSON 对象"""
        if not text:
            return None
        try:
            text = text.strip()
            if '```' in text:
                text = text.split('```json')[-1].split('```')[0]
                text = text.strip()
            if not text.startswith('{'):
                idx = text.find('{')
                if idx >= 0:
                    text = text[idx:]
            return json.loads(text)
        except:
            return None


alignment_service = AlignmentService()
