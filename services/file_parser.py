"""
文件内容提取工具 - 支持 Word, PDF, Excel, TXT 等格式
"""
import os
import re
import shutil
import subprocess
import tempfile


def extract_text_from_file(filepath: str) -> str:
    """
    从文件中提取纯文本内容。
    支持: .txt, .pdf, .docx, .doc, .xlsx, .xls, .csv, .md
    """
    ext = os.path.splitext(filepath)[1].lower()

    try:
        if ext in ('.txt', '.md', '.csv', '.log'):
            return _read_text_file(filepath)
        elif ext == '.pdf':
            return _read_pdf(filepath)
        elif ext in ('.docx',):
            return _read_docx(filepath)
        elif ext in ('.doc',):
            return _read_doc(filepath)
        elif ext in ('.xlsx', '.xls'):
            return _read_excel(filepath)
        else:
            return f"[不支持的文件格式: {ext}]"
    except Exception as e:
        return f"[文件解析失败: {str(e)}]"


def _read_text_file(filepath: str) -> str:
    """读取纯文本文件，自动检测编码"""
    try:
        import chardet
        with open(filepath, 'rb') as f:
            raw = f.read()
        detected = chardet.detect(raw)
        encoding = detected.get('encoding', 'utf-8') or 'utf-8'
        return raw.decode(encoding, errors='replace')
    except ImportError:
        # chardet 不可用时尝试常见编码
        for enc in ('utf-8', 'gbk', 'gb2312', 'latin-1'):
            try:
                with open(filepath, 'r', encoding=enc) as f:
                    return f.read()
            except (UnicodeDecodeError, LookupError):
                continue
        return "[文本文件编码无法识别]"


def _read_pdf(filepath: str) -> str:
    """读取 PDF 文件"""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(filepath)
        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                pages.append(f"--- 第{i+1}页 ---\n{text.strip()}")
        return "\n\n".join(pages) if pages else "[PDF 无可提取文本（可能是扫描件）]"
    except ImportError:
        return "[需要安装 PyPDF2: pip install PyPDF2]"


def _read_docx(filepath: str) -> str:
    """读取 Word (.docx) 文件"""
    try:
        from docx import Document
        doc = Document(filepath)
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # 也读取表格
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                paragraphs.append("\n[表格]\n" + "\n".join(rows))

        return "\n".join(paragraphs) if paragraphs else "[Word 文档内容为空]"
    except ImportError:
        return "[需要安装 python-docx: pip install python-docx]"


def _read_doc(filepath: str) -> str:
    """
    读取 Word (.doc) 文件。
    优先尝试系统能力（win32com / antiword / catdoc / soffice），
    不可用时使用二进制启发式提取可读中文行。
    """
    readers = (
        _read_doc_via_win32com,
        _read_doc_via_antiword,
        _read_doc_via_catdoc,
        _read_doc_via_soffice_convert,
        _read_doc_binary_heuristic
    )
    for reader in readers:
        try:
            text = reader(filepath)
            if _is_meaningful_doc_text(text):
                return text
        except Exception:
            continue
    return "[DOC 文件解析失败：当前环境缺少可用解析能力]"


def _is_meaningful_doc_text(text: str) -> bool:
    if not text:
        return False
    value = str(text).strip()
    if not value or value.startswith('['):
        return False
    if len(value) < 20:
        return False
    return bool(re.search(r'[\u4e00-\u9fffA-Za-z]{8,}', value))


def _read_doc_via_win32com(filepath: str) -> str:
    try:
        import pythoncom
        import win32com.client  # type: ignore
    except Exception:
        return ""

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.txt')
    tmp_path = tmp.name
    tmp.close()
    word = None
    doc = None
    try:
        pythoncom.CoInitialize()
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(os.path.abspath(filepath), ReadOnly=True)
        # 7 = wdFormatText
        doc.SaveAs(os.path.abspath(tmp_path), FileFormat=7)
        return _read_text_file(tmp_path)
    except Exception:
        return ""
    finally:
        try:
            if doc is not None:
                doc.Close(False)
        except Exception:
            pass
        try:
            if word is not None:
                word.Quit()
        except Exception:
            pass
        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass
        try:
            os.remove(tmp_path)
        except Exception:
            pass


def _read_doc_via_antiword(filepath: str) -> str:
    if not shutil.which('antiword'):
        return ""
    try:
        result = subprocess.run(
            ['antiword', filepath],
            capture_output=True,
            text=True,
            encoding='utf-8',
            errors='ignore',
            timeout=25
        )
        if result.returncode == 0:
            return result.stdout.strip()
    except Exception:
        pass
    return ""


def _read_doc_via_catdoc(filepath: str) -> str:
    if not shutil.which('catdoc'):
        return ""
    try:
        result = subprocess.run(
            ['catdoc', filepath],
            capture_output=True,
            text=True,
            encoding='utf-8',
            errors='ignore',
            timeout=25
        )
        if result.returncode == 0:
            return result.stdout.strip()
    except Exception:
        pass
    return ""


def _read_doc_via_soffice_convert(filepath: str) -> str:
    cmd = shutil.which('soffice') or shutil.which('libreoffice')
    if not cmd:
        return ""

    tmp_dir = tempfile.mkdtemp(prefix='doc2txt_')
    out_name = os.path.splitext(os.path.basename(filepath))[0] + '.txt'
    out_path = os.path.join(tmp_dir, out_name)
    try:
        result = subprocess.run(
            [cmd, '--headless', '--convert-to', 'txt:Text', '--outdir', tmp_dir, filepath],
            capture_output=True,
            text=True,
            encoding='utf-8',
            errors='ignore',
            timeout=45
        )
        if result.returncode != 0 or not os.path.exists(out_path):
            return ""
        return _read_text_file(out_path)
    except Exception:
        return ""
    finally:
        try:
            if os.path.exists(out_path):
                os.remove(out_path)
            os.rmdir(tmp_dir)
        except Exception:
            pass


_DOC_MEDICAL_KEYWORDS = [
    '评估', '护理', '记录', '压疮', '压力性损伤', '损伤', '措施', '分值', '分级',
    '科室', '床号', '姓名', '日期', '签名', '责任护士', '护士长', '评分', '风险',
    '院内', '带入', '不可分期', '住院号', '入院日期', '出院', '死亡', '好转',
    '治愈', '恶化', '诊断', '部位', '口是', '口否'
]


def _doc_line_score(line: str) -> float:
    txt = str(line or '').strip()
    if len(txt) < 2 or len(txt) > 180:
        return -1.0

    if re.search(r'[^\u4e00-\u9fffA-Za-z0-9\s\-\+\(\)（）\[\]【】:：;；,，.。/|%℃<>~=_#]', txt):
        return -1.0

    cjk_count = len(re.findall(r'[\u4e00-\u9fff]', txt))
    if cjk_count == 0:
        return -1.0

    score = cjk_count / max(1, len(txt))

    if any(key in txt for key in _DOC_MEDICAL_KEYWORDS):
        score += 2.1

    if re.search(r'(^\d+期[:：]|口[\u4e00-\u9fff]|[一二三四五六七八九十]+期[:：])', txt):
        score += 0.5

    if re.search(r'[：:|口□]', txt):
        score += 0.3

    if re.search(r'\d', txt):
        score += 0.2

    if len(txt) <= 4:
        score -= 0.2

    return score


def _read_doc_binary_heuristic(filepath: str) -> str:
    """
    legacy .doc 二进制回退提取：
    对 utf-16le / gb18030 解码后的文本做行级评分，筛出可读医疗表单语句。
    """
    try:
        payload = open(filepath, 'rb').read()
    except Exception:
        return ""

    candidates = {}
    order = []

    for enc in ('utf-16le', 'gb18030'):
        try:
            decoded = payload.decode(enc, errors='ignore')
        except Exception:
            continue

        raw_lines = re.split(r'[\r\n\x00\t]+', decoded)
        for raw in raw_lines:
            for seg in re.split(r'(?<=。)|(?<=；)|[ ]{2,}', raw):
                text = seg.strip()
                if not text:
                    continue
                score = _doc_line_score(text)
                if score < 1.5:
                    continue
                if text not in candidates:
                    order.append(text)
                    candidates[text] = score
                elif score > candidates[text]:
                    candidates[text] = score

    if not candidates:
        return ""

    lines = [line for line in order if candidates.get(line, 0) >= 1.5]
    return '\n'.join(lines[:400]).strip()


def _read_excel(filepath: str) -> str:
    """读取 Excel 文件"""
    try:
        from openpyxl import load_workbook
        wb = load_workbook(filepath, read_only=True, data_only=True)
        sheets_text = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else '' for c in row]
                if any(cells):  # 跳过全空行
                    rows.append(" | ".join(cells))
            if rows:
                sheets_text.append(f"[工作表: {sheet_name}]\n" + "\n".join(rows))

        wb.close()
        return "\n\n".join(sheets_text) if sheets_text else "[Excel 文件内容为空]"
    except ImportError:
        return "[需要安装 openpyxl: pip install openpyxl]"


# 支持的文件扩展名
SUPPORTED_EXTENSIONS = {'.txt', '.md', '.csv', '.log', '.pdf', '.docx', '.doc', '.xlsx', '.xls'}


def is_supported(filename: str) -> bool:
    """检查文件是否支持解析"""
    ext = os.path.splitext(filename)[1].lower()
    return ext in SUPPORTED_EXTENSIONS
