import json
import os
import re
import statistics
from datetime import date, datetime, timedelta
from typing import Any, Dict, List, Optional

from ai_utils import call_ai
from database import DatabasePool
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from services.audit_service import audit_service


RND_TARGET_ROLE_HINTS = ('研发', '开发', '后端', '前端', '测试', '产品', '架构', '算法', '平台', '接口研发')
ONSITE_REVIEWER_ROLE_HINTS = ('现场', '驻场', '实施', '交付', '项目经理', '工程', 'PM')
EXCLUDED_ROLE_HINTS = ('销售', '商务', '财务', '行政', '采购', '助理', '客户', '院方')
RECOGNITION_WEIGHTS = {
    'gratitude': 6,
    'support': 8,
    'mentoring': 7,
    'rescue': 10,
    'customer_praise': 10,
}


class PerformanceReviewService:
    @staticmethod
    def _parse_date(value: Any) -> Optional[date]:
        if value is None or value == '':
            return None
        if isinstance(value, date) and not isinstance(value, datetime):
            return value
        if isinstance(value, datetime):
            return value.date()
        text = str(value).strip()
        if not text:
            return None
        for fmt in ('%Y-%m-%d', '%Y-%m-%d %H:%M:%S', '%Y-%m-%dT%H:%M:%S'):
            try:
                return datetime.strptime(text[:19], fmt).date()
            except ValueError:
                continue
        return None

    @staticmethod
    def _parse_datetime(value: Any) -> Optional[datetime]:
        if value is None or value == '':
            return None
        if isinstance(value, datetime):
            return value
        if isinstance(value, date):
            return datetime.combine(value, datetime.min.time())
        text = str(value).strip()
        if not text:
            return None
        normalized = text.replace('T', ' ')[:19]
        for fmt in ('%Y-%m-%d %H:%M:%S', '%Y-%m-%d'):
            try:
                return datetime.strptime(normalized, fmt)
            except ValueError:
                continue
        return None

    @staticmethod
    def _clamp(value: float, lower: float = 0.0, upper: float = 100.0) -> float:
        return max(lower, min(upper, float(value)))

    @staticmethod
    def _round(value: float, digits: int = 2) -> float:
        return round(float(value), digits)

    @staticmethod
    def _current_cycle_range(ref: Optional[date] = None):
        current = ref or date.today()
        monday = current - timedelta(days=current.weekday())
        sunday = monday + timedelta(days=6)
        iso_year, iso_week, _ = monday.isocalendar()
        cycle_key = f"{iso_year}-W{iso_week:02d}"
        title = f"{iso_year}年第{iso_week:02d}周研发绩效周评"
        return cycle_key, title, monday, sunday

    @staticmethod
    def _date_window_bounds(start_date: str, end_date: str):
        start_day = PerformanceReviewService._parse_date(start_date)
        end_day = PerformanceReviewService._parse_date(end_date)
        if not start_day or not end_day:
            return f"{start_date} 00:00:00", f"{end_date} 23:59:59", end_date
        start_dt = f"{start_day.isoformat()} 00:00:00"
        next_day = end_day + timedelta(days=1)
        end_exclusive = f"{next_day.isoformat()} 00:00:00"
        return start_dt, end_exclusive, end_day.isoformat()

    @staticmethod
    def _trimmed_mean(values: List[float]) -> Optional[float]:
        cleaned = [float(v) for v in values if v is not None]
        if not cleaned:
            return None
        if len(cleaned) >= 3:
            cleaned = sorted(cleaned)[1:-1]
        return sum(cleaned) / max(1, len(cleaned))

    @staticmethod
    def _as_bool(value: Any) -> bool:
        if isinstance(value, bool):
            return value
        return str(value).strip().lower() in ('1', 'true', 'yes', 'y', 'on')

    @staticmethod
    def _json_load(value: Any, fallback):
        if not value:
            return fallback
        if isinstance(value, (dict, list)):
            return value
        try:
            return json.loads(value)
        except Exception:
            return fallback

    @staticmethod
    def _extract_json_array(text: str) -> List[Dict[str, Any]]:
        if not text:
            return []
        fenced = re.findall(r'```json\s*([\s\S]*?)```', text, flags=re.IGNORECASE)
        candidates = fenced + [text]
        for candidate in candidates:
            start = candidate.find('[')
            end = candidate.rfind(']')
            if start == -1 or end == -1 or end <= start:
                continue
            try:
                payload = json.loads(candidate[start:end + 1])
                return payload if isinstance(payload, list) else []
            except Exception:
                continue
        return []

    @staticmethod
    def _is_rnd_target_member(member: Dict[str, Any], signals: Dict[str, int]) -> bool:
        role = str(member.get('role') or '')
        name = str(member.get('name') or '')
        has_include_hint = any(keyword in role or keyword in name for keyword in RND_TARGET_ROLE_HINTS)
        has_exclude_hint = any(keyword in role for keyword in EXCLUDED_ROLE_HINTS)
        if has_exclude_hint:
            return False
        if has_include_hint:
            return True
        # Fallback: if the member is not an onsite implementation reviewer, treat them as an R&D-side target.
        return not PerformanceReviewService._is_onsite_reviewer(member)

    @staticmethod
    def _is_onsite_reviewer(member: Dict[str, Any]) -> bool:
        role = str(member.get('role') or '')
        name = str(member.get('name') or '')
        if any(keyword in role for keyword in EXCLUDED_ROLE_HINTS):
            return False
        if any(keyword in role or keyword in name for keyword in RND_TARGET_ROLE_HINTS):
            return False
        return PerformanceReviewService._as_bool(member.get('is_onsite')) or any(keyword in role for keyword in ONSITE_REVIEWER_ROLE_HINTS)

    @staticmethod
    def _classify_project_member_identity(member: Dict[str, Any]) -> str:
        role = str(member.get('role') or '')
        if any(keyword in role for keyword in EXCLUDED_ROLE_HINTS):
            return 'unknown'
        if PerformanceReviewService._is_onsite_reviewer(member):
            return 'implementation'
        if PerformanceReviewService._is_rnd_target_member(member, {}):
            return 'rnd'
        return 'unknown'

    @staticmethod
    def ensure_cycle(ref_date: Optional[str] = None, operator: Optional[str] = None) -> Dict[str, Any]:
        parsed_ref = PerformanceReviewService._parse_date(ref_date) if ref_date else None
        cycle_key, title, start_date, end_date = PerformanceReviewService._current_cycle_range(parsed_ref)

        with DatabasePool.get_connection() as conn:
            row = conn.execute(
                DatabasePool.format_sql('SELECT * FROM performance_review_cycles WHERE cycle_key = ?'),
                (cycle_key,),
            ).fetchone()
            if row:
                cycle = dict(row)
            else:
                sql = '''
                    INSERT INTO performance_review_cycles
                    (cycle_key, title, start_date, end_date, status, created_by, created_at, updated_at)
                    VALUES (?, ?, ?, ?, ?, ?, CURRENT_TIMESTAMP, CURRENT_TIMESTAMP)
                '''
                conn.execute(
                    DatabasePool.format_sql(sql),
                    (cycle_key, title, start_date.isoformat(), end_date.isoformat(), 'active', operator or '系统'),
                )
                conn.commit()
                cycle = dict(conn.execute(
                    DatabasePool.format_sql('SELECT * FROM performance_review_cycles WHERE cycle_key = ?'),
                    (cycle_key,),
                ).fetchone())
        return cycle

    @staticmethod
    def list_cycles(limit: int = 16) -> List[Dict[str, Any]]:
        with DatabasePool.get_connection() as conn:
            rows = conn.execute(
                DatabasePool.format_sql('''
                    SELECT * FROM performance_review_cycles
                    ORDER BY start_date DESC
                    LIMIT ?
                '''),
                (limit,),
            ).fetchall()
            return [dict(row) for row in rows]

    @staticmethod
    def _get_cycle(conn, cycle_id: Optional[int] = None, ref_date: Optional[str] = None):
        if cycle_id:
            row = conn.execute(
                DatabasePool.format_sql('SELECT * FROM performance_review_cycles WHERE id = ?'),
                (cycle_id,),
            ).fetchone()
            if row:
                return dict(row)
        return PerformanceReviewService.ensure_cycle(ref_date=ref_date)

    @staticmethod
    def _sync_targets_for_cycle(conn, cycle: Dict[str, Any], project_id: Optional[int] = None) -> List[Dict[str, Any]]:
        cycle_id = cycle['id']
        start_date = str(cycle['start_date'])
        end_date = str(cycle['end_date'])
        start_dt, end_exclusive_dt, _ = PerformanceReviewService._date_window_bounds(start_date, end_date)

        project_filter = 'AND pm.project_id = ?' if project_id else ''
        active_members = conn.execute(DatabasePool.format_sql(f'''
            SELECT
                pm.*,
                p.project_name,
                p.hospital_name,
                p.status AS project_status
            FROM project_members pm
            JOIN projects p ON p.id = pm.project_id
            WHERE pm.status = '在岗'
              AND p.status != '已删除'
              {project_filter}
            ORDER BY p.project_name, pm.name
        '''), (project_id,) if project_id else ()).fetchall()

        if project_id:
            conn.execute(
                DatabasePool.format_sql('UPDATE performance_review_targets SET status = ? WHERE cycle_id = ? AND project_id = ?'),
                ('inactive', cycle_id, project_id),
            )
        else:
            conn.execute(
                DatabasePool.format_sql('UPDATE performance_review_targets SET status = ? WHERE cycle_id = ?'),
                ('inactive', cycle_id),
            )

        for member in active_members:
            member_dict = dict(member)
            log_signal = conn.execute(DatabasePool.format_sql('''
                SELECT COUNT(*) AS c
                FROM work_logs
                WHERE project_id = ?
                  AND log_date BETWEEN ? AND ?
                  AND (
                    member_id = ?
                    OR member_name = ?
                  )
            '''), (
                member_dict['project_id'],
                start_date,
                end_date,
                member_dict['id'],
                member_dict['name'],
            )).fetchone()['c'] or 0

            task_signal = conn.execute(DatabasePool.format_sql('''
                SELECT COUNT(*) AS c
                FROM tasks t
                JOIN project_stages s ON s.id = t.stage_id
                WHERE s.project_id = ?
                  AND COALESCE(t.assigned_to, '') = ?
                  AND (
                    (t.completed_date BETWEEN ? AND ?)
                    OR (t.updated_at IS NOT NULL AND t.updated_at >= ? AND t.updated_at < ?)
                    OR t.is_completed = 0
                  )
            '''), (
                member_dict['project_id'],
                member_dict['name'],
                start_date,
                end_date,
                start_dt,
                end_exclusive_dt,
            )).fetchone()['c'] or 0

            stage_signal = conn.execute(DatabasePool.format_sql('''
                SELECT COUNT(*) AS c
                FROM project_stages
                WHERE project_id = ?
                  AND COALESCE(responsible_person, '') = ?
            '''), (
                member_dict['project_id'],
                member_dict['name'],
            )).fetchone()['c'] or 0

            issue_signal = conn.execute(DatabasePool.format_sql('''
                SELECT COUNT(*) AS c
                FROM issues
                WHERE project_id = ?
                  AND owner_member_id = ?
                  AND created_at < ?
                  AND (
                    resolved_at IS NULL
                    OR resolved_at >= ?
                  )
            '''), (
                member_dict['project_id'],
                member_dict['id'],
                end_exclusive_dt,
                start_dt,
            )).fetchone()['c'] or 0

            signals = {
                'logs': int(log_signal),
                'tasks': int(task_signal),
                'stages': int(stage_signal),
                'issues': int(issue_signal),
            }
            if not PerformanceReviewService._is_rnd_target_member(member_dict, signals):
                continue

            conn.execute(DatabasePool.format_sql('''
                INSERT INTO performance_review_targets (
                    cycle_id, project_id, member_id, member_name, member_role, is_onsite, status,
                    signal_logs, signal_tasks, signal_stages, signal_issues, created_at, updated_at
                )
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, CURRENT_TIMESTAMP, CURRENT_TIMESTAMP)
                ON CONFLICT(cycle_id, project_id, member_name)
                DO UPDATE SET
                    member_id = excluded.member_id,
                    member_role = excluded.member_role,
                    is_onsite = excluded.is_onsite,
                    status = excluded.status,
                    signal_logs = excluded.signal_logs,
                    signal_tasks = excluded.signal_tasks,
                    signal_stages = excluded.signal_stages,
                    signal_issues = excluded.signal_issues,
                    updated_at = CURRENT_TIMESTAMP
            '''), (
                cycle_id,
                member_dict['project_id'],
                member_dict['id'],
                member_dict['name'],
                member_dict.get('role') or '',
                bool(PerformanceReviewService._as_bool(member_dict.get('is_onsite'))),
                'active',
                signals['logs'],
                signals['tasks'],
                signals['stages'],
                signals['issues'],
            ))

        conn.commit()
        rows_filter = 'AND t.project_id = ?' if project_id else ''
        rows = conn.execute(DatabasePool.format_sql(f'''
            SELECT
                t.*,
                p.project_name,
                p.hospital_name
            FROM performance_review_targets t
            JOIN projects p ON p.id = t.project_id
            WHERE t.cycle_id = ?
              AND t.status = 'active'
              {rows_filter}
            ORDER BY p.project_name, t.member_name
        '''), (cycle_id, project_id) if project_id else (cycle_id,)).fetchall()
        return [dict(row) for row in rows]

    @staticmethod
    def _form_total(form: Dict[str, Any]) -> float:
        return (
            float(form.get('score_responsibility') or 0) * 0.30 +
            float(form.get('score_collaboration') or 0) * 0.25 +
            float(form.get('score_response') or 0) * 0.25 +
            float(form.get('score_professional') or 0) * 0.20
        )

    @staticmethod
    def _build_fallback_ai(target: Dict[str, Any], onsite_score: float, evidence_score: float, warmth_score: float,
                           evidence: Dict[str, Any], review_count: int) -> Dict[str, Any]:
        active_days = evidence.get('worklog', {}).get('active_days', 0)
        completed_tasks = evidence.get('delivery', {}).get('completed_tasks', 0)
        closed_issues = evidence.get('issues', {}).get('closed_owned_issues', 0)

        if review_count >= 2 and onsite_score >= 88:
            highlight = '现场协作评价很稳，周内执行可信度高。'
        elif completed_tasks >= 2 or closed_issues >= 2:
            highlight = '本周闭环动作扎实，交付推进比较有节奏。'
        elif active_days >= 4:
            highlight = '投入连续性不错，现场存在感和跟进度都在线。'
        else:
            highlight = '本周证据量偏少，但整体交付节奏保持平稳。'

        risk = '问题归属和日志沉淀还可以再结构化一些，能让评分更客观。'
        if evidence.get('issues', {}).get('high_pending_owned', 0) > 0:
            risk = '仍有高优先级问题挂在名下，建议优先清掉再进入新任务。'
        elif review_count == 0:
            risk = '当前缺少现场同事评分，建议尽快补齐，以免结果过度依赖系统证据。'

        support = '建议保留一段固定时间做复盘沉淀，把解决方案写进日志或知识库。'
        if warmth_score >= 90:
            support = '建议把这周的支援经验沉淀成一条最佳实践，顺手带一下团队。'
        elif onsite_score < 75:
            support = '建议下周主动同步关键进展，减少协作方对信息透明度的担忧。'

        base = evidence_score * 0.75 + onsite_score * 0.15 + warmth_score * 0.10
        ai_score = PerformanceReviewService._clamp(base)
        summary = f"AI 认为本周综合表现为 {PerformanceReviewService._round(ai_score, 1)} 分，亮点在于{highlight.replace('。', '')}。"
        return {
            'ai_score': PerformanceReviewService._round(ai_score, 2),
            'highlight': highlight,
            'risk': risk,
            'support': support,
            'summary': summary,
        }

    @staticmethod
    def _build_score_payload(conn, cycle: Dict[str, Any], target: Dict[str, Any]) -> Dict[str, Any]:
        cycle_id = cycle['id']
        start_date = str(cycle['start_date'])
        end_date = str(cycle['end_date'])
        start_dt, end_exclusive_dt, _ = PerformanceReviewService._date_window_bounds(start_date, end_date)
        target_id = target['id']
        project_id = target['project_id']
        member_id = target.get('member_id')
        member_name = target['member_name']

        forms = conn.execute(DatabasePool.format_sql('''
            SELECT * FROM performance_review_forms
            WHERE cycle_id = ? AND target_id = ?
            ORDER BY created_at DESC
        '''), (cycle_id, target_id)).fetchall()
        forms = [dict(row) for row in forms]

        recognitions = conn.execute(DatabasePool.format_sql('''
            SELECT * FROM performance_recognition
            WHERE cycle_id = ? AND target_id = ?
            ORDER BY created_at DESC
        '''), (cycle_id, target_id)).fetchall()
        recognitions = [dict(row) for row in recognitions]

        worklogs = conn.execute(DatabasePool.format_sql('''
            SELECT *
            FROM work_logs
            WHERE project_id = ?
              AND log_date BETWEEN ? AND ?
              AND (
                member_id = ?
                OR member_name = ?
              )
            ORDER BY log_date DESC
        '''), (project_id, start_date, end_date, member_id, member_name)).fetchall()
        worklogs = [dict(row) for row in worklogs]

        completed_tasks = conn.execute(DatabasePool.format_sql('''
            SELECT COUNT(*) AS c
            FROM tasks t
            JOIN project_stages s ON s.id = t.stage_id
            WHERE s.project_id = ?
              AND COALESCE(t.assigned_to, '') = ?
              AND t.is_completed = 1
              AND t.completed_date BETWEEN ? AND ?
        '''), (project_id, member_name, start_date, end_date)).fetchone()['c'] or 0

        active_tasks = conn.execute(DatabasePool.format_sql('''
            SELECT COUNT(*) AS c
            FROM tasks t
            JOIN project_stages s ON s.id = t.stage_id
            WHERE s.project_id = ?
              AND COALESCE(t.assigned_to, '') = ?
        '''), (project_id, member_name)).fetchone()['c'] or 0

        completed_stages = conn.execute(DatabasePool.format_sql('''
            SELECT COUNT(*) AS c
            FROM project_stages
            WHERE project_id = ?
              AND COALESCE(responsible_person, '') = ?
              AND actual_end_date BETWEEN ? AND ?
        '''), (project_id, member_name, start_date, end_date)).fetchone()['c'] or 0

        owned_issues = conn.execute(DatabasePool.format_sql('''
            SELECT *
            FROM issues
            WHERE project_id = ?
              AND owner_member_id = ?
              AND created_at < ?
              AND (
                resolved_at IS NULL
                OR resolved_at >= ?
              )
            ORDER BY created_at DESC
        '''), (project_id, member_id, end_exclusive_dt, start_dt)).fetchall()
        owned_issues = [dict(row) for row in owned_issues]

        project_satisfaction = conn.execute(DatabasePool.format_sql('''
            SELECT AVG(score_overall) AS avg_overall, COUNT(*) AS count
            FROM customer_satisfaction
            WHERE project_id = ?
              AND survey_date BETWEEN ? AND ?
        '''), (project_id, start_date, end_date)).fetchone()

        onsite_dimension_values = {
            'responsibility': [form.get('score_responsibility') for form in forms],
            'collaboration': [form.get('score_collaboration') for form in forms],
            'response': [form.get('score_response') for form in forms],
            'professional': [form.get('score_professional') for form in forms],
        }
        onsite_dimensions = {
            key: PerformanceReviewService._round(
                PerformanceReviewService._trimmed_mean([value for value in values if value is not None]) or 0,
                2,
            )
            for key, values in onsite_dimension_values.items()
        }
        onsite_total = PerformanceReviewService._trimmed_mean([PerformanceReviewService._form_total(form) for form in forms])
        onsite_review_count = len(forms)
        onsite_neutral_used = False
        if onsite_total is None:
            onsite_total = 78.0
            onsite_neutral_used = True
        onsite_total = PerformanceReviewService._round(onsite_total, 2)

        active_days = len({str(log.get('log_date')) for log in worklogs if log.get('log_date')})
        worklog_completeness_ratios = []
        for log in worklogs:
            filled = 0
            filled += 1 if str(log.get('work_content') or '').strip() else 0
            filled += 1 if str(log.get('tomorrow_plan') or '').strip() else 0
            filled += 1 if str(log.get('issues_encountered') or '').strip() else 0
            worklog_completeness_ratios.append(filled / 3)
        worklog_completeness = statistics.mean(worklog_completeness_ratios) if worklog_completeness_ratios else 0

        delivery_continuity_points = PerformanceReviewService._clamp(active_days / 5 * 15, 0, 15)
        task_delivery_points = min(15.0, completed_tasks * 5.0)
        closure_contribution_points = min(10.0, completed_stages * 5.0)

        issue_created_in_cycle = [
            item for item in owned_issues
            if start_date <= str(item.get('created_at', ''))[:10] <= end_date
        ]
        responded_in_24h = 0
        for issue in issue_created_in_cycle:
            created_at = PerformanceReviewService._parse_datetime(issue.get('created_at'))
            first_response_at = PerformanceReviewService._parse_datetime(issue.get('first_response_at'))
            if created_at and first_response_at and (first_response_at - created_at).total_seconds() <= 24 * 3600:
                responded_in_24h += 1
        response_rate = (responded_in_24h / len(issue_created_in_cycle)) if issue_created_in_cycle else None

        closed_owned_issues = sum(
            1
            for issue in owned_issues
            if issue.get('resolved_at') and start_date <= str(issue.get('resolved_at'))[:10] <= end_date
        )
        issue_window_total = len(owned_issues)
        closure_rate = (closed_owned_issues / issue_window_total) if issue_window_total else None
        reopen_count = sum(int(issue.get('reopen_count') or 0) for issue in owned_issues)
        high_pending_owned = sum(
            1
            for issue in owned_issues
            if issue.get('severity') == '高' and issue.get('status') not in ('已解决', '已关闭')
        )

        if response_rate is None:
            response_points = 9.0
            response_note = '本周期没有结构化归属到该成员的新问题，响应时效按中性分计入。'
        else:
            response_points = response_rate * 15.0
            response_note = f'24h 内首次响应率 {responded_in_24h}/{len(issue_created_in_cycle)}。'

        if closure_rate is None:
            closure_points = 8.0
            closure_note = '本周期缺少归属问题闭环样本，闭环率按中性分计入。'
        else:
            closure_points = closure_rate * 10.0
            closure_note = f'归属问题闭环 {closed_owned_issues}/{issue_window_total}。'

        discipline_points = max(0.0, 5.0 - min(5.0, reopen_count * 1.5 + high_pending_owned * 1.2))
        log_attendance_points = PerformanceReviewService._clamp(active_days / 5 * 12.0, 0, 12.0)
        log_quality_points = PerformanceReviewService._clamp(worklog_completeness * 8.0, 0, 8.0)

        satisfaction_avg = float(project_satisfaction['avg_overall'] or 0) if project_satisfaction else 0
        satisfaction_points = (satisfaction_avg / 5.0) * 6.0 if satisfaction_avg else 4.0
        blocker_points = 3.0 if not owned_issues else max(0.0, 4.0 - min(4.0, high_pending_owned))

        ai_evidence_score = sum([
            delivery_continuity_points,
            task_delivery_points,
            closure_contribution_points,
            response_points,
            closure_points,
            discipline_points,
            log_attendance_points,
            log_quality_points,
            satisfaction_points,
            blocker_points,
        ])
        ai_evidence_score = PerformanceReviewService._round(ai_evidence_score, 2)

        recognition_points = 0
        for recognition in recognitions:
            recognition_points += RECOGNITION_WEIGHTS.get(recognition.get('recognition_type'), 6)
        review_highlight_points = min(10.0, sum(1 for form in forms if str(form.get('highlight') or '').strip()) * 3.0)
        customer_positive_points = 0.0
        if satisfaction_avg >= 4.6:
            customer_positive_points = 10.0
        elif satisfaction_avg >= 4.0:
            customer_positive_points = 6.0
        warmth_score = min(100.0, 70.0 + min(20.0, recognition_points) + review_highlight_points + customer_positive_points)
        warmth_score = PerformanceReviewService._round(warmth_score, 2)

        highlights = [str(form.get('highlight') or '').strip() for form in forms if str(form.get('highlight') or '').strip()]
        suggestions = [str(form.get('suggestion') or '').strip() for form in forms if str(form.get('suggestion') or '').strip()]
        evidence_notes = [str(form.get('evidence_note') or '').strip() for form in forms if str(form.get('evidence_note') or '').strip()]

        formula = {
            'final_formula': '最终分 = 现场评价分×35% + AI证据分×55% + 人情味分×10% + 人工校准',
            'onsite_formula': '现场评价分 = 去极值后均值(责任心30% + 协作度25% + 响应速度25% + 专业度20%)',
            'ai_formula': {
                'delivery_execution_40': {
                    '连续投入_15': PerformanceReviewService._round(delivery_continuity_points, 2),
                    '任务兑现_15': PerformanceReviewService._round(task_delivery_points, 2),
                    '里程碑闭环贡献_10': PerformanceReviewService._round(closure_contribution_points, 2),
                },
                'issue_closure_30': {
                    '响应时效_15': PerformanceReviewService._round(response_points, 2),
                    '闭环率_10': PerformanceReviewService._round(closure_points, 2),
                    '纪律性_5': PerformanceReviewService._round(discipline_points, 2),
                },
                'process_transparency_20': {
                    '日志出勤_12': PerformanceReviewService._round(log_attendance_points, 2),
                    '日志完整度_8': PerformanceReviewService._round(log_quality_points, 2),
                },
                'quality_stability_10': {
                    '客户认可_6': PerformanceReviewService._round(satisfaction_points, 2),
                    '高优稳定性_4': PerformanceReviewService._round(blocker_points, 2),
                },
            },
            'warmth_formula': {
                '基础信任分_70': 70,
                '暖心卡奖励': min(20.0, recognition_points),
                '现场亮点反馈': PerformanceReviewService._round(review_highlight_points, 2),
                '客户正反馈': PerformanceReviewService._round(customer_positive_points, 2),
            },
            'neutral_rules': {
                'onsite': '若暂无现场评分，现场评价分按 78 分中性值临时计入。',
                'issue_response': response_note,
                'issue_closure': closure_note,
            },
        }

        evidence = {
            'project': {
                'project_id': project_id,
                'project_name': target.get('project_name'),
                'hospital_name': target.get('hospital_name'),
            },
            'signals': {
                'worklogs': len(worklogs),
                'active_days': active_days,
                'completed_tasks': completed_tasks,
                'active_tasks': active_tasks,
                'completed_stages': completed_stages,
                'owned_issues': len(owned_issues),
                'recognitions': len(recognitions),
                'review_count': onsite_review_count,
            },
            'worklog': {
                'active_days': active_days,
                'completeness_ratio': PerformanceReviewService._round(worklog_completeness, 4),
                'latest_logs': [
                    {
                        'log_date': log.get('log_date'),
                        'work_content': str(log.get('work_content') or '')[:120],
                        'tomorrow_plan': str(log.get('tomorrow_plan') or '')[:120],
                    }
                    for log in worklogs[:3]
                ],
            },
            'delivery': {
                'completed_tasks': completed_tasks,
                'active_tasks': active_tasks,
                'completed_stages': completed_stages,
            },
            'issues': {
                'owned_issues': len(owned_issues),
                'created_in_cycle': len(issue_created_in_cycle),
                'responded_in_24h': responded_in_24h,
                'closed_owned_issues': closed_owned_issues,
                'reopen_count': reopen_count,
                'high_pending_owned': high_pending_owned,
            },
            'reviews': {
                'count': onsite_review_count,
                'highlights': highlights[:5],
                'suggestions': suggestions[:5],
                'evidence_notes': evidence_notes[:5],
            },
            'recognitions': recognitions[:5],
            'customer_feedback': {
                'avg_overall': PerformanceReviewService._round(satisfaction_avg, 2),
                'count': int(project_satisfaction['count'] or 0) if project_satisfaction else 0,
            },
        }

        fallback_ai = PerformanceReviewService._build_fallback_ai(
            target,
            onsite_total,
            ai_evidence_score,
            warmth_score,
            evidence,
            onsite_review_count,
        )

        return {
            'target_id': target_id,
            'project_id': project_id,
            'target_member_name': member_name,
            'onsite_score': onsite_total,
            'onsite_dimensions': onsite_dimensions,
            'onsite_neutral_used': onsite_neutral_used,
            'review_count': onsite_review_count,
            'ai_evidence_score': ai_evidence_score,
            'warmth_score': warmth_score,
            'formula': formula,
            'evidence': evidence,
            'fallback_ai': fallback_ai,
            'highlights': highlights[:5],
            'suggestions': suggestions[:5],
        }

    @staticmethod
    def _save_scorecard(conn, cycle: Dict[str, Any], target: Dict[str, Any], payload: Dict[str, Any],
                        ai_map: Optional[Dict[int, Dict[str, Any]]] = None):
        existing = conn.execute(DatabasePool.format_sql('''
            SELECT * FROM performance_score_cards
            WHERE cycle_id = ? AND target_id = ?
        '''), (cycle['id'], target['id'])).fetchone()
        existing = dict(existing) if existing else {}

        ai_fallback = payload['fallback_ai']
        ai_payload = (ai_map or {}).get(target['id'], {})
        ai_raw_score = float(ai_payload.get('ai_score') or ai_fallback['ai_score'] or payload['ai_evidence_score'])
        ai_score = PerformanceReviewService._clamp(
            ai_raw_score,
            payload['ai_evidence_score'] - 8,
            payload['ai_evidence_score'] + 8,
        )
        ai_score = PerformanceReviewService._round(ai_score, 2)
        calibration_delta = float(existing.get('calibration_delta') or 0)
        base_final_score = (
            float(payload['onsite_score']) * 0.35 +
            float(ai_score) * 0.55 +
            float(payload['warmth_score']) * 0.10
        )
        base_final_score = PerformanceReviewService._round(base_final_score, 2)
        final_score = PerformanceReviewService._round(
            PerformanceReviewService._clamp(base_final_score + calibration_delta),
            2,
        )

        ai_summary = ai_payload.get('summary') or ai_fallback['summary']
        ai_highlight = ai_payload.get('highlight') or ai_fallback['highlight']
        ai_risk = ai_payload.get('risk') or ai_fallback['risk']
        ai_support = ai_payload.get('support') or ai_fallback['support']

        conn.execute(DatabasePool.format_sql('''
            INSERT INTO performance_score_cards (
                cycle_id, target_id, project_id, target_member_name, onsite_score, ai_evidence_score,
                ai_raw_score, ai_score, warmth_score, base_final_score, calibration_delta, final_score,
                formula_json, evidence_json, ai_summary, ai_highlight, ai_risk, ai_support,
                review_count, calibrated_by, calibrated_reason, approved_by, approved_at,
                ai_generated_at, generated_at, updated_at
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, CURRENT_TIMESTAMP, CURRENT_TIMESTAMP)
            ON CONFLICT(cycle_id, target_id)
            DO UPDATE SET
                project_id = excluded.project_id,
                target_member_name = excluded.target_member_name,
                onsite_score = excluded.onsite_score,
                ai_evidence_score = excluded.ai_evidence_score,
                ai_raw_score = excluded.ai_raw_score,
                ai_score = excluded.ai_score,
                warmth_score = excluded.warmth_score,
                base_final_score = excluded.base_final_score,
                final_score = excluded.final_score,
                formula_json = excluded.formula_json,
                evidence_json = excluded.evidence_json,
                ai_summary = excluded.ai_summary,
                ai_highlight = excluded.ai_highlight,
                ai_risk = excluded.ai_risk,
                ai_support = excluded.ai_support,
                review_count = excluded.review_count,
                calibrated_by = ?,
                calibrated_reason = ?,
                approved_by = ?,
                approved_at = ?,
                ai_generated_at = excluded.ai_generated_at,
                updated_at = CURRENT_TIMESTAMP
        '''), (
            cycle['id'],
            target['id'],
            target['project_id'],
            target['member_name'],
            payload['onsite_score'],
            payload['ai_evidence_score'],
            ai_raw_score,
            ai_score,
            payload['warmth_score'],
            base_final_score,
            calibration_delta,
            final_score,
            json.dumps(payload['formula'], ensure_ascii=False),
            json.dumps(payload['evidence'], ensure_ascii=False),
            ai_summary,
            ai_highlight,
            ai_risk,
            ai_support,
            payload['review_count'],
            existing.get('calibrated_by'),
            existing.get('calibrated_reason'),
            existing.get('approved_by'),
            existing.get('approved_at'),
            datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            existing.get('calibrated_by'),
            existing.get('calibrated_reason'),
            existing.get('approved_by'),
            existing.get('approved_at'),
        ))

    @staticmethod
    def _refresh_ai_guidance(conn, cycle: Dict[str, Any], targets: List[Dict[str, Any]], payloads: Dict[int, Dict[str, Any]]):
        prompt_members = []
        for target in targets:
            payload = payloads[target['id']]
            evidence = payload['evidence']
            prompt_members.append({
                'target_id': target['id'],
                'member_name': target['member_name'],
                'member_role': target.get('member_role') or '',
                'project_name': target.get('project_name') or '',
                'onsite_score': payload['onsite_score'],
                'review_count': payload['review_count'],
                'ai_evidence_score': payload['ai_evidence_score'],
                'warmth_score': payload['warmth_score'],
                'active_days': evidence['worklog']['active_days'],
                'completed_tasks': evidence['delivery']['completed_tasks'],
                'completed_stages': evidence['delivery']['completed_stages'],
                'closed_owned_issues': evidence['issues']['closed_owned_issues'],
                'high_pending_owned': evidence['issues']['high_pending_owned'],
                'highlights': payload['highlights'][:2],
                'suggestions': payload['suggestions'][:2],
            })

        system_prompt = """你是一位公平、公正但有人情味的研发绩效周评教练。
你的任务不是随意打分，而是基于系统已有证据给出谨慎的 AI 建议分。
要求：
1. 输出 JSON 数组，不要输出额外说明。
2. 每项包含 target_id, ai_score, highlight, risk, support, summary。
3. ai_score 为 0-100 的整数或小数，语气克制，不夸张。
4. highlight 要指出可被认可的地方。
5. risk 要指出最值得关注的风险。
6. support 要给出有温度、可执行的支持建议。
7. summary 用 1-2 句话概括，避免官话。
"""

        user_prompt = "请基于以下本周成员证据生成绩效 AI 建议分：\n" + json.dumps(prompt_members, ensure_ascii=False)
        text = call_ai(user_prompt, task_type='summary', system_prompt=system_prompt)
        rows = PerformanceReviewService._extract_json_array(text)
        ai_map = {}
        for row in rows:
            if not isinstance(row, dict) or not row.get('target_id'):
                continue
            target_id = int(row['target_id'])
            ai_map[target_id] = {
                'ai_score': PerformanceReviewService._clamp(row.get('ai_score') or 0),
                'highlight': str(row.get('highlight') or '').strip(),
                'risk': str(row.get('risk') or '').strip(),
                'support': str(row.get('support') or '').strip(),
                'summary': str(row.get('summary') or '').strip(),
            }
        return ai_map

    @staticmethod
    def rebuild_cycle(cycle_id: int, use_ai: bool = True, operator: Optional[str] = None, project_id: Optional[int] = None) -> Dict[str, Any]:
        with DatabasePool.get_connection() as conn:
            cycle = PerformanceReviewService._get_cycle(conn, cycle_id=cycle_id)
            if str(cycle.get('status') or '') == 'locked':
                raise ValueError('当前周期已锁定，不能重算分数')
            targets = PerformanceReviewService._sync_targets_for_cycle(conn, cycle, project_id=project_id)
            payloads = {target['id']: PerformanceReviewService._build_score_payload(conn, cycle, target) for target in targets}

            ai_map = {}
            if use_ai and targets:
                try:
                    ai_map = PerformanceReviewService._refresh_ai_guidance(conn, cycle, targets, payloads)
                except Exception:
                    ai_map = {}

            for target in targets:
                PerformanceReviewService._save_scorecard(conn, cycle, target, payloads[target['id']], ai_map=ai_map)

            conn.execute(
                DatabasePool.format_sql('UPDATE performance_review_cycles SET updated_at = CURRENT_TIMESTAMP WHERE id = ?'),
                (cycle['id'],),
            )
            conn.commit()

        if operator:
            audit_service.log_operation(operator, '重算绩效周评', 'performance_cycle', cycle_id, cycle.get('title'))

        return {
            'cycle_id': cycle_id,
            'project_id': project_id,
            'targets': len(targets),
            'ai_used': bool(ai_map),
        }

    @staticmethod
    def _current_user_can_review_project(conn, user: Dict[str, Any], project_id: int) -> bool:
        reviewer_name = str(user.get('display_name') or user.get('username') or '').strip()
        if not reviewer_name:
            return False
        row = conn.execute(DatabasePool.format_sql('''
            SELECT *
            FROM project_members
            WHERE project_id = ?
              AND status = '在岗'
              AND name = ?
        '''), (project_id, reviewer_name)).fetchone()
        return PerformanceReviewService._is_onsite_reviewer(dict(row)) if row else False

    @staticmethod
    def _current_user_matches_member(user: Dict[str, Any], member_name: str) -> bool:
        actor = str(user.get('display_name') or user.get('username') or '').strip()
        return bool(actor and actor == str(member_name or '').strip())

    @staticmethod
    def _get_project_role_readiness(conn, project_id: int) -> Dict[str, Any]:
        rows = conn.execute(DatabasePool.format_sql('''
            SELECT id, name, role, is_onsite, current_city, status
            FROM project_members
            WHERE project_id = ?
              AND status = '在岗'
            ORDER BY role, name
        '''), (project_id,)).fetchall()
        members = [dict(row) for row in rows]
        implementation = []
        rnd = []
        unknown = []
        for member in members:
            identity = PerformanceReviewService._classify_project_member_identity(member)
            payload = {
                'id': member.get('id'),
                'name': member.get('name') or '',
                'role': member.get('role') or '',
                'current_city': member.get('current_city') or '',
                'is_onsite': bool(member.get('is_onsite')),
            }
            if identity == 'implementation':
                implementation.append(payload)
            elif identity == 'rnd':
                rnd.append(payload)
            else:
                unknown.append(payload)
        return {
            'implementation_reviewers': implementation,
            'rnd_targets': rnd,
            'unknown_members': unknown,
            'ready': bool(implementation and rnd and not unknown),
        }

    @staticmethod
    def _resolve_user_ids_by_names(conn, names: List[str]) -> List[int]:
        cleaned = []
        for name in names or []:
            value = str(name or '').strip()
            if value and value not in cleaned:
                cleaned.append(value)
        if not cleaned:
            return []
        placeholders = ','.join(['?' for _ in cleaned])
        rows = conn.execute(DatabasePool.format_sql(f'''
            SELECT id, username, display_name
            FROM users
            WHERE username IN ({placeholders}) OR display_name IN ({placeholders})
        '''), tuple(cleaned + cleaned)).fetchall()
        return list({int(row['id']) for row in rows if row and row.get('id')})

    @staticmethod
    def _create_targeted_notifications(conn, project_id: int, user_ids: List[int], title: str, content: str,
                                       notification_type: str = 'info', dedupe_hours: int = 12) -> int:
        unique_user_ids = []
        for uid in user_ids or []:
            try:
                value = int(uid)
            except Exception:
                continue
            if value not in unique_user_ids:
                unique_user_ids.append(value)
        if not unique_user_ids:
            return 0

        created = 0
        cutoff = (datetime.now() - timedelta(hours=max(1, int(dedupe_hours or 12)))).strftime('%Y-%m-%d %H:%M:%S')
        for uid in unique_user_ids:
            existing = conn.execute(DatabasePool.format_sql('''
                SELECT id
                FROM notifications
                WHERE project_id = ?
                  AND target_user_id = ?
                  AND title = ?
                  AND content = ?
                  AND created_at >= ?
                ORDER BY id DESC
                LIMIT 1
            '''), (project_id, uid, title, content, cutoff)).fetchone()
            if existing:
                continue
            conn.execute(DatabasePool.format_sql('''
                INSERT INTO notifications (project_id, target_user_id, title, content, type, remind_type)
                VALUES (?, ?, ?, ?, ?, ?)
            '''), (project_id, uid, title, content, notification_type, 'once'))
            created += 1
        return created

    @staticmethod
    def _send_project_review_reminders_internal(conn, cycle_id: int, project_id: int, operator: str = '系统') -> Dict[str, Any]:
        cycle = conn.execute(DatabasePool.format_sql('SELECT * FROM performance_review_cycles WHERE id = ?'), (cycle_id,)).fetchone()
        if not cycle:
            raise ValueError('未找到周评周期')
        cycle = dict(cycle)
        project = conn.execute(DatabasePool.format_sql('SELECT project_name FROM projects WHERE id = ?'), (project_id,)).fetchone()
        project_name = (dict(project) if project else {}).get('project_name') or f'项目#{project_id}'
        readiness = PerformanceReviewService._get_project_role_readiness(conn, project_id)

        targets = conn.execute(DatabasePool.format_sql('''
            SELECT t.member_name, COALESCE(s.review_count, 0) AS review_count
            FROM performance_review_targets t
            LEFT JOIN performance_score_cards s
                ON s.cycle_id = t.cycle_id AND s.target_id = t.id
            WHERE t.cycle_id = ?
              AND t.project_id = ?
              AND t.status = 'active'
            ORDER BY t.member_name
        '''), (cycle_id, project_id)).fetchall()
        pending_targets = [dict(row) for row in targets if int(row.get('review_count') or 0) < 2]
        if not pending_targets:
            return {'created': 0, 'message': '当前项目没有待提醒的研发评分对象'}

        reviewer_names = [item.get('name') for item in readiness.get('implementation_reviewers') or []]
        reviewer_user_ids = PerformanceReviewService._resolve_user_ids_by_names(conn, reviewer_names)
        if not reviewer_user_ids:
            return {'created': 0, 'message': '未匹配到实施侧评分人的系统账号，无法定向提醒'}

        target_names = '、'.join(item.get('member_name') for item in pending_targets[:8])
        title = f"🏅 待完成研发绩效评分：{project_name}"
        content = f"{cycle.get('title') or '本周'}还有 {len(pending_targets)} 位研发对象待实施侧评分：{target_names}。请进入绩效页完成评价。"
        created = PerformanceReviewService._create_targeted_notifications(
            conn,
            project_id,
            reviewer_user_ids,
            title,
            content,
            'warning',
            dedupe_hours=4
        )
        return {'created': created, 'message': f'已发送 {created} 条提醒', 'project_name': project_name}

    @staticmethod
    def _ensure_cycle_editable(conn, cycle_id: int):
        row = conn.execute(
            DatabasePool.format_sql('SELECT status FROM performance_review_cycles WHERE id = ?'),
            (cycle_id,),
        ).fetchone()
        status = str((dict(row) if row else {}).get('status') or 'active')
        if status == 'locked':
            raise ValueError('当前周期已锁定，不能再修改评分内容')

    @staticmethod
    def get_overview(cycle_id: Optional[int] = None, ref_date: Optional[str] = None, project_id: Optional[int] = None,
                     current_user: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        current_user = current_user or {}
        with DatabasePool.get_connection() as conn:
            cycle = PerformanceReviewService._get_cycle(conn, cycle_id=cycle_id, ref_date=ref_date)
            scorecard_meta = conn.execute(DatabasePool.format_sql('''
                SELECT COUNT(*) AS c, MIN(ai_generated_at) AS earliest_ai_generated_at
                FROM performance_score_cards
                WHERE cycle_id = ?
                  AND (? IS NULL OR project_id = ?)
            '''), (cycle['id'], project_id, project_id)).fetchone()
        cycle_locked = str(cycle.get('status')) == 'locked'
        should_refresh_ai = bool(
            (not cycle_locked) and str(cycle.get('status')) == 'active' and (
                not scorecard_meta
                or not int(scorecard_meta['c'] or 0)
                or not scorecard_meta['earliest_ai_generated_at']
            )
        )
        should_rebuild = bool(
            (not cycle_locked) and (
                should_refresh_ai
                or not scorecard_meta
                or not int(scorecard_meta['c'] or 0)
            )
        )
        if should_rebuild:
            PerformanceReviewService.rebuild_cycle(cycle['id'], use_ai=should_refresh_ai, project_id=project_id)

        with DatabasePool.get_connection() as conn:
            project_clause = 'AND t.project_id = ?' if project_id else ''
            cycle_rows = conn.execute(DatabasePool.format_sql(f'''
                SELECT
                    t.*,
                    p.project_name,
                    p.hospital_name,
                    s.id AS scorecard_id,
                    s.onsite_score,
                    s.ai_evidence_score,
                    s.ai_raw_score,
                    s.ai_score,
                    s.warmth_score,
                    s.base_final_score,
                    s.calibration_delta,
                    s.final_score,
                    s.formula_json,
                    s.evidence_json,
                    s.ai_summary,
                    s.ai_highlight,
                    s.ai_risk,
                    s.ai_support,
                    s.review_count,
                    s.calibrated_by,
                    s.calibrated_reason,
                    s.approved_by,
                    s.approved_at,
                    s.updated_at
                FROM performance_review_targets t
                JOIN projects p ON p.id = t.project_id
                LEFT JOIN performance_score_cards s
                    ON s.cycle_id = t.cycle_id AND s.target_id = t.id
                WHERE t.cycle_id = ?
                  AND t.status = 'active'
                  {project_clause}
                ORDER BY p.project_name, COALESCE(s.final_score, 0) DESC, t.member_name
            '''), (cycle['id'], project_id) if project_id else (cycle['id'],)).fetchall()

            projects_map: Dict[int, Dict[str, Any]] = {}
            all_members = []
            warm_wall = []
            pending_appeals = []
            reviewer_name = str(current_user.get('display_name') or current_user.get('username') or '').strip()
            cycle_locked = str(cycle.get('status') or '') == 'locked'
            for row in cycle_rows:
                item = dict(row)
                project_id = item['project_id']
                project_can_review = PerformanceReviewService._current_user_can_review_project(conn, current_user, project_id)
                member_can_appeal = PerformanceReviewService._current_user_matches_member(current_user, item['member_name'])
                formula = PerformanceReviewService._json_load(item.get('formula_json'), {})
                evidence = PerformanceReviewService._json_load(item.get('evidence_json'), {})
                recognitions = conn.execute(DatabasePool.format_sql('''
                    SELECT * FROM performance_recognition
                    WHERE cycle_id = ? AND target_id = ?
                    ORDER BY created_at DESC
                    LIMIT 4
                '''), (cycle['id'], item['id'])).fetchall()
                recognitions = [dict(rec) for rec in recognitions]
                my_form = conn.execute(DatabasePool.format_sql('''
                    SELECT * FROM performance_review_forms
                    WHERE cycle_id = ? AND target_id = ? AND reviewer_name = ?
                    ORDER BY updated_at DESC LIMIT 1
                '''), (cycle['id'], item['id'], reviewer_name)).fetchone()
                my_form = dict(my_form) if my_form else None
                appeals = conn.execute(DatabasePool.format_sql('''
                    SELECT * FROM performance_appeals
                    WHERE cycle_id = ? AND target_id = ?
                    ORDER BY created_at DESC
                '''), (cycle['id'], item['id'])).fetchall()
                appeals = [dict(appeal) for appeal in appeals]
                latest_pending_appeal = next((appeal for appeal in appeals if appeal.get('status') == 'pending'), None)

                member_payload = {
                    'target_id': item['id'],
                    'scorecard_id': item.get('scorecard_id'),
                    'member_id': item.get('member_id'),
                    'member_name': item['member_name'],
                    'member_role': item.get('member_role') or '',
                    'is_onsite': bool(item.get('is_onsite')),
                    'signals': {
                        'logs': int(item.get('signal_logs') or 0),
                        'tasks': int(item.get('signal_tasks') or 0),
                        'stages': int(item.get('signal_stages') or 0),
                        'issues': int(item.get('signal_issues') or 0),
                    },
                    'scores': {
                        'onsite': float(item.get('onsite_score') or 0),
                        'ai_evidence': float(item.get('ai_evidence_score') or 0),
                        'ai_raw': float(item.get('ai_raw_score') or 0),
                        'ai': float(item.get('ai_score') or 0),
                        'warmth': float(item.get('warmth_score') or 0),
                        'base_final': float(item.get('base_final_score') or 0),
                        'calibration_delta': float(item.get('calibration_delta') or 0),
                        'final': float(item.get('final_score') or 0),
                    },
                    'review_count': int(item.get('review_count') or 0),
                    'formula': formula,
                    'evidence': evidence,
                    'ai_summary': item.get('ai_summary') or '',
                    'ai_highlight': item.get('ai_highlight') or '',
                    'ai_risk': item.get('ai_risk') or '',
                    'ai_support': item.get('ai_support') or '',
                    'calibrated_by': item.get('calibrated_by') or '',
                    'calibrated_reason': item.get('calibrated_reason') or '',
                    'approved_by': item.get('approved_by') or '',
                    'approved_at': item.get('approved_at'),
                    'recognitions': recognitions,
                    'appeals': appeals,
                    'my_review': my_form,
                    'actions': {
                        'can_review': bool((not cycle_locked) and project_can_review and reviewer_name and reviewer_name != item['member_name']),
                        'can_recognize': bool((not cycle_locked) and project_can_review and reviewer_name and reviewer_name != item['member_name']),
                        'can_calibrate': bool((not cycle_locked) and str(current_user.get('role') or '') in ('admin', 'project_manager')),
                        'can_approve': str(current_user.get('role') or '') in ('admin', 'project_manager'),
                        'can_appeal': bool(member_can_appeal and item.get('scorecard_id')),
                        'can_resolve_appeal': str(current_user.get('role') or '') in ('admin', 'project_manager'),
                    },
                }
                all_members.append(member_payload)
                if latest_pending_appeal:
                    pending_appeals.append({
                        'appeal_id': latest_pending_appeal['id'],
                        'scorecard_id': item.get('scorecard_id'),
                        'member_name': item['member_name'],
                        'project_name': item['project_name'],
                        'appellant_name': latest_pending_appeal.get('appellant_name') or '',
                        'appeal_reason': latest_pending_appeal.get('appeal_reason') or '',
                        'created_at': latest_pending_appeal.get('created_at'),
                    })
                warm_wall.extend([
                    {
                        'member_name': item['member_name'],
                        'project_name': item['project_name'],
                        'title': rec.get('title') or '感谢你',
                        'content': rec.get('content') or '',
                        'recognition_type': rec.get('recognition_type') or '',
                        'giver_name': rec.get('giver_name') or '',
                        'created_at': rec.get('created_at'),
                    }
                    for rec in recognitions[:2]
                ])

                project_entry = projects_map.setdefault(project_id, {
                    'project_id': project_id,
                    'project_name': item['project_name'],
                    'hospital_name': item['hospital_name'],
                    'members': [],
                })
                project_entry['members'].append(member_payload)

            avg_final = statistics.mean([member['scores']['final'] for member in all_members]) if all_members else 0
            avg_onsite = statistics.mean([member['scores']['onsite'] for member in all_members]) if all_members else 0
            avg_ai = statistics.mean([member['scores']['ai'] for member in all_members]) if all_members else 0
            pending_reviews = sum(1 for member in all_members if member['review_count'] < 2)
            reviewed_targets = sum(1 for member in all_members if member['review_count'] >= 2)
            approved_targets = sum(1 for member in all_members if member.get('approved_by'))

            leaderboard = sorted(all_members, key=lambda item: item['scores']['final'], reverse=True)[:10]
            warm_rank = sorted(all_members, key=lambda item: item['scores']['warmth'], reverse=True)[:10]
            readiness = PerformanceReviewService._get_project_role_readiness(conn, project_id) if project_id else None

            return {
                'cycle': cycle,
                'project_id': project_id,
                'cycles': PerformanceReviewService.list_cycles(),
                'summary': {
                    'member_count': len(all_members),
                    'project_count': len(projects_map),
                    'avg_final_score': PerformanceReviewService._round(avg_final, 2),
                    'avg_onsite_score': PerformanceReviewService._round(avg_onsite, 2),
                    'avg_ai_score': PerformanceReviewService._round(avg_ai, 2),
                    'pending_reviews': pending_reviews,
                    'reviewed_targets': reviewed_targets,
                    'approved_targets': approved_targets,
                    'recognition_count': len(warm_wall),
                    'pending_appeals': len(pending_appeals),
                },
                'projects': list(projects_map.values()),
                'leaderboard': leaderboard,
                'warm_rankings': warm_rank,
                'warm_wall': warm_wall[:12],
                'pending_appeals': pending_appeals[:12],
                'readiness': readiness,
                'permissions': {
                    'can_calibrate': str(current_user.get('role') or '') in ('admin', 'project_manager'),
                    'can_approve': str(current_user.get('role') or '') in ('admin', 'project_manager'),
                    'can_resolve_appeal': str(current_user.get('role') or '') in ('admin', 'project_manager'),
                },
            }

    @staticmethod
    def save_review_form(cycle_id: int, target_id: int, data: Dict[str, Any], current_user: Dict[str, Any]) -> Dict[str, Any]:
        reviewer_name = str(current_user.get('display_name') or current_user.get('username') or '').strip()
        if not reviewer_name:
            raise ValueError('无法识别当前评分人')
        project_id = None

        with DatabasePool.get_connection() as conn:
            PerformanceReviewService._ensure_cycle_editable(conn, cycle_id)
            target = conn.execute(DatabasePool.format_sql('''
                SELECT * FROM performance_review_targets WHERE id = ? AND cycle_id = ?
            '''), (target_id, cycle_id)).fetchone()
            if not target:
                raise ValueError('未找到评分对象')
            target = dict(target)
            if reviewer_name == target['member_name']:
                raise ValueError('不能给自己评分')
            if not PerformanceReviewService._current_user_can_review_project(conn, current_user, target['project_id']):
                raise ValueError('当前账号没有该项目的现场评分资格')
            project_id = target['project_id']

            responsibility = PerformanceReviewService._clamp(data.get('score_responsibility') or 0)
            collaboration = PerformanceReviewService._clamp(data.get('score_collaboration') or 0)
            response = PerformanceReviewService._clamp(data.get('score_response') or 0)
            professional = PerformanceReviewService._clamp(data.get('score_professional') or 0)

            conn.execute(DatabasePool.format_sql('''
                INSERT INTO performance_review_forms (
                    cycle_id, target_id, project_id, reviewer_name, reviewer_role,
                    score_responsibility, score_collaboration, score_response, score_professional,
                    highlight, suggestion, evidence_note, created_at, updated_at
                )
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, CURRENT_TIMESTAMP, CURRENT_TIMESTAMP)
                ON CONFLICT(cycle_id, target_id, reviewer_name)
                DO UPDATE SET
                    reviewer_role = excluded.reviewer_role,
                    score_responsibility = excluded.score_responsibility,
                    score_collaboration = excluded.score_collaboration,
                    score_response = excluded.score_response,
                    score_professional = excluded.score_professional,
                    highlight = excluded.highlight,
                    suggestion = excluded.suggestion,
                    evidence_note = excluded.evidence_note,
                    updated_at = CURRENT_TIMESTAMP
            '''), (
                cycle_id,
                target_id,
                target['project_id'],
                reviewer_name,
                current_user.get('role') or '',
                responsibility,
                collaboration,
                response,
                professional,
                str(data.get('highlight') or '').strip(),
                str(data.get('suggestion') or '').strip(),
                str(data.get('evidence_note') or '').strip(),
            ))
            conn.commit()

        PerformanceReviewService.rebuild_cycle(cycle_id, use_ai=True, operator=reviewer_name, project_id=project_id)
        with DatabasePool.get_connection() as conn:
            try:
                PerformanceReviewService._send_project_review_reminders_internal(conn, cycle_id, project_id, operator=reviewer_name)
                conn.commit()
            except Exception:
                pass
        audit_service.log_operation(reviewer_name, '提交现场周评', 'performance_target', target_id, target['member_name'], new_val=data)
        return {'success': True}

    @staticmethod
    def add_recognition(cycle_id: int, target_id: int, data: Dict[str, Any], current_user: Dict[str, Any]) -> Dict[str, Any]:
        giver_name = str(current_user.get('display_name') or current_user.get('username') or '').strip()
        if not giver_name:
            raise ValueError('无法识别当前用户')

        with DatabasePool.get_connection() as conn:
            PerformanceReviewService._ensure_cycle_editable(conn, cycle_id)
            target = conn.execute(DatabasePool.format_sql('''
                SELECT * FROM performance_review_targets WHERE id = ? AND cycle_id = ?
            '''), (target_id, cycle_id)).fetchone()
            if not target:
                raise ValueError('未找到被认可对象')
            target = dict(target)
            if giver_name == target['member_name']:
                raise ValueError('不能给自己发暖心卡')
            if not PerformanceReviewService._current_user_can_review_project(conn, current_user, target['project_id']):
                raise ValueError('当前账号没有该项目的暖心加分资格')

            recognition_type = str(data.get('recognition_type') or 'gratitude').strip() or 'gratitude'
            conn.execute(DatabasePool.format_sql('''
                INSERT INTO performance_recognition (
                    cycle_id, target_id, project_id, giver_name, recognition_type,
                    title, content, weight, created_at
                )
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, CURRENT_TIMESTAMP)
            '''), (
                cycle_id,
                target_id,
                target['project_id'],
                giver_name,
                recognition_type,
                str(data.get('title') or '谢谢你').strip(),
                str(data.get('content') or '').strip(),
                RECOGNITION_WEIGHTS.get(recognition_type, 6),
            ))
            conn.commit()

        PerformanceReviewService.rebuild_cycle(cycle_id, use_ai=True, operator=giver_name)
        audit_service.log_operation(giver_name, '发送暖心卡', 'performance_target', target_id, target['member_name'], new_val=data)
        return {'success': True}

    @staticmethod
    def calibrate_scorecard(scorecard_id: int, data: Dict[str, Any], current_user: Dict[str, Any]) -> Dict[str, Any]:
        role = str(current_user.get('role') or '')
        operator = str(current_user.get('display_name') or current_user.get('username') or '').strip() or '系统'
        if role not in ('admin', 'project_manager'):
            raise ValueError('当前账号没有人工校准权限')

        delta = float(data.get('delta') or 0)
        if delta < -5 or delta > 5:
            raise ValueError('人工校准范围必须在 -5 到 5 分之间')
        reason = str(data.get('reason') or '').strip()
        if not reason:
            raise ValueError('请填写人工校准理由')

        with DatabasePool.get_connection() as conn:
            row = conn.execute(DatabasePool.format_sql('''
                SELECT * FROM performance_score_cards WHERE id = ?
            '''), (scorecard_id,)).fetchone()
            if not row:
                raise ValueError('未找到评分卡')
            scorecard = dict(row)
            PerformanceReviewService._ensure_cycle_editable(conn, scorecard['cycle_id'])
            final_score = PerformanceReviewService._round(
                PerformanceReviewService._clamp(float(scorecard.get('base_final_score') or 0) + delta),
                2,
            )
            conn.execute(DatabasePool.format_sql('''
                INSERT INTO performance_adjustments (
                    scorecard_id, cycle_id, target_id, operator, delta, reason, created_at
                )
                VALUES (?, ?, ?, ?, ?, ?, CURRENT_TIMESTAMP)
            '''), (
                scorecard_id,
                scorecard['cycle_id'],
                scorecard['target_id'],
                operator,
                delta,
                reason,
            ))
            conn.execute(DatabasePool.format_sql('''
                UPDATE performance_score_cards
                SET calibration_delta = ?,
                    final_score = ?,
                    calibrated_by = ?,
                    calibrated_reason = ?,
                    approved_by = NULL,
                    approved_at = NULL,
                    updated_at = CURRENT_TIMESTAMP
                WHERE id = ?
            '''), (
                delta,
                final_score,
                operator,
                reason,
                scorecard_id,
            ))
            conn.commit()

        audit_service.log_operation(operator, '人工校准绩效分', 'performance_scorecard', scorecard_id, scorecard.get('target_member_name'), old_val={'delta': scorecard.get('calibration_delta')}, new_val={'delta': delta, 'reason': reason})
        return {'success': True}

    @staticmethod
    def approve_scorecard(scorecard_id: int, current_user: Dict[str, Any]) -> Dict[str, Any]:
        role = str(current_user.get('role') or '')
        operator = str(current_user.get('display_name') or current_user.get('username') or '').strip() or '系统'
        if role not in ('admin', 'project_manager'):
            raise ValueError('当前账号没有审批权限')

        with DatabasePool.get_connection() as conn:
            row = conn.execute(DatabasePool.format_sql('SELECT * FROM performance_score_cards WHERE id = ?'), (scorecard_id,)).fetchone()
            if not row:
                raise ValueError('未找到评分卡')
            scorecard = dict(row)
            target_user_ids = PerformanceReviewService._resolve_user_ids_by_names(conn, [scorecard.get('target_member_name')])
            conn.execute(DatabasePool.format_sql('''
                UPDATE performance_score_cards
                SET approved_by = ?, approved_at = CURRENT_TIMESTAMP, updated_at = CURRENT_TIMESTAMP
                WHERE id = ?
            '''), (operator, scorecard_id))
            PerformanceReviewService._create_targeted_notifications(
                conn,
                scorecard.get('project_id'),
                target_user_ids,
                '🏅 研发绩效评分已审批',
                f"你在当前项目中的研发绩效评分已由 {operator} 审批，可进入绩效页查看最终结果。",
                'info',
                dedupe_hours=6
            )
            conn.commit()

        audit_service.log_operation(operator, '审批绩效评分卡', 'performance_scorecard', scorecard_id, scorecard.get('target_member_name'))
        return {'success': True}

    @staticmethod
    def get_member_trend(project_id: int, member_name: str, limit: int = 12) -> Dict[str, Any]:
        with DatabasePool.get_connection() as conn:
            rows = conn.execute(DatabasePool.format_sql('''
                SELECT
                    c.id AS cycle_id,
                    c.title,
                    c.start_date,
                    c.end_date,
                    s.id AS scorecard_id,
                    s.onsite_score,
                    s.ai_score,
                    s.ai_raw_score,
                    s.warmth_score,
                    s.final_score,
                    s.calibration_delta,
                    s.review_count,
                    s.approved_by,
                    (
                        SELECT COUNT(*)
                        FROM performance_appeals a
                        WHERE a.scorecard_id = s.id
                    ) AS appeal_count
                FROM performance_score_cards s
                JOIN performance_review_cycles c ON c.id = s.cycle_id
                WHERE s.project_id = ?
                  AND s.target_member_name = ?
                ORDER BY c.start_date DESC
                LIMIT ?
            '''), (project_id, member_name, limit)).fetchall()
            items = [dict(row) for row in rows]

        trend = list(reversed(items))
        avg_final = statistics.mean([float(item.get('final_score') or 0) for item in trend]) if trend else 0
        best_final = max([float(item.get('final_score') or 0) for item in trend], default=0)
        lowest_final = min([float(item.get('final_score') or 0) for item in trend], default=0)
        return {
            'project_id': project_id,
            'member_name': member_name,
            'trend': trend,
            'summary': {
                'cycles': len(trend),
                'avg_final_score': PerformanceReviewService._round(avg_final, 2),
                'best_final_score': PerformanceReviewService._round(best_final, 2),
                'lowest_final_score': PerformanceReviewService._round(lowest_final, 2),
            },
        }

    @staticmethod
    def submit_appeal(scorecard_id: int, data: Dict[str, Any], current_user: Dict[str, Any]) -> Dict[str, Any]:
        appellant_name = str(current_user.get('display_name') or current_user.get('username') or '').strip()
        if not appellant_name:
            raise ValueError('无法识别当前申诉人')
        appeal_reason = str(data.get('appeal_reason') or '').strip()
        if not appeal_reason:
            raise ValueError('请填写申诉理由')

        with DatabasePool.get_connection() as conn:
            row = conn.execute(DatabasePool.format_sql('SELECT * FROM performance_score_cards WHERE id = ?'), (scorecard_id,)).fetchone()
            if not row:
                raise ValueError('未找到评分卡')
            scorecard = dict(row)
            if not PerformanceReviewService._current_user_matches_member(current_user, scorecard.get('target_member_name')):
                raise ValueError('当前账号只能对自己的评分卡发起申诉')

            existing = conn.execute(DatabasePool.format_sql('''
                SELECT id FROM performance_appeals
                WHERE scorecard_id = ? AND status = 'pending'
                ORDER BY created_at DESC LIMIT 1
            '''), (scorecard_id,)).fetchone()
            if existing:
                raise ValueError('当前评分卡已有待处理申诉')

            pm_users = []
            project = conn.execute(DatabasePool.format_sql('SELECT project_manager, project_name FROM projects WHERE id = ?'), (scorecard['project_id'],)).fetchone()
            project = dict(project) if project else {}
            if project.get('project_manager'):
                pm_users = PerformanceReviewService._resolve_user_ids_by_names(conn, [project.get('project_manager')])
            admin_rows = conn.execute(DatabasePool.format_sql("SELECT id FROM users WHERE role = 'admin'")).fetchall()
            admin_ids = [int(row['id']) for row in admin_rows if row and row.get('id')]

            conn.execute(DatabasePool.format_sql('''
                INSERT INTO performance_appeals (
                    scorecard_id, cycle_id, target_id, appellant_name, appeal_reason,
                    status, created_at, updated_at
                )
                VALUES (?, ?, ?, ?, ?, 'pending', CURRENT_TIMESTAMP, CURRENT_TIMESTAMP)
            '''), (
                scorecard_id,
                scorecard['cycle_id'],
                scorecard['target_id'],
                appellant_name,
                appeal_reason,
            ))
            PerformanceReviewService._create_targeted_notifications(
                conn,
                scorecard.get('project_id'),
                list({*pm_users, *admin_ids}),
                '⚠️ 研发绩效有待处理申诉',
                f"{appellant_name} 对当前项目的研发绩效评分发起了申诉，请及时复核处理。",
                'warning',
                dedupe_hours=4
            )
            conn.commit()

        audit_service.log_operation(appellant_name, '发起绩效申诉', 'performance_scorecard', scorecard_id, scorecard.get('target_member_name'), new_val={'appeal_reason': appeal_reason})
        return {'success': True}

    @staticmethod
    def resolve_appeal(appeal_id: int, data: Dict[str, Any], current_user: Dict[str, Any]) -> Dict[str, Any]:
        role = str(current_user.get('role') or '')
        resolver = str(current_user.get('display_name') or current_user.get('username') or '').strip() or '系统'
        if role not in ('admin', 'project_manager'):
            raise ValueError('当前账号没有复核权限')

        resolution_text = str(data.get('resolution_text') or '').strip()
        if not resolution_text:
            raise ValueError('请填写复核结论')

        with DatabasePool.get_connection() as conn:
            row = conn.execute(DatabasePool.format_sql('SELECT * FROM performance_appeals WHERE id = ?'), (appeal_id,)).fetchone()
            if not row:
                raise ValueError('未找到申诉记录')
            appeal = dict(row)
            scorecard = conn.execute(DatabasePool.format_sql('SELECT project_id FROM performance_score_cards WHERE id = ?'), (appeal.get('scorecard_id'),)).fetchone()
            scorecard = dict(scorecard) if scorecard else {}
            target_user_ids = PerformanceReviewService._resolve_user_ids_by_names(conn, [appeal.get('appellant_name')])
            conn.execute(DatabasePool.format_sql('''
                UPDATE performance_appeals
                SET status = ?,
                    resolution_text = ?,
                    resolved_by = ?,
                    resolved_at = CURRENT_TIMESTAMP,
                    updated_at = CURRENT_TIMESTAMP
                WHERE id = ?
            '''), (
                data.get('status') or 'resolved',
                resolution_text,
                resolver,
                appeal_id,
            ))
            PerformanceReviewService._create_targeted_notifications(
                conn,
                scorecard.get('project_id'),
                target_user_ids,
                '✅ 研发绩效申诉已处理',
                f"你的绩效申诉已由 {resolver} 处理：{resolution_text[:120]}",
                'info',
                dedupe_hours=4
            )
            conn.commit()

        audit_service.log_operation(resolver, '处理绩效申诉', 'performance_appeal', appeal_id, appeal.get('appellant_name'), new_val={'status': data.get('status') or 'resolved', 'resolution_text': resolution_text})
        return {'success': True}

    @staticmethod
    def get_team_trend(project_id: Optional[int] = None, limit: int = 12) -> Dict[str, Any]:
        with DatabasePool.get_connection() as conn:
            project_filter = 'WHERE (? IS NULL OR s.project_id = ?)' 
            rows = conn.execute(DatabasePool.format_sql(f'''
                SELECT
                    c.id AS cycle_id,
                    c.title,
                    c.start_date,
                    c.end_date,
                    c.status,
                    COUNT(s.id) AS member_count,
                    AVG(s.final_score) AS avg_final_score,
                    AVG(s.onsite_score) AS avg_onsite_score,
                    AVG(s.ai_score) AS avg_ai_score,
                    AVG(s.warmth_score) AS avg_warmth_score,
                    SUM(CASE WHEN s.approved_by IS NOT NULL THEN 1 ELSE 0 END) AS approved_count,
                    (
                        SELECT COUNT(*)
                        FROM performance_appeals a
                        WHERE a.cycle_id = c.id AND a.status = 'pending'
                    ) AS pending_appeals
                FROM performance_review_cycles c
                LEFT JOIN performance_score_cards s ON s.cycle_id = c.id
                {project_filter}
                GROUP BY c.id, c.title, c.start_date, c.end_date, c.status
                ORDER BY c.start_date DESC
                LIMIT ?
            '''), (project_id, project_id, limit)).fetchall()
            items = [dict(row) for row in rows]

        trend = list(reversed(items))
        return {
            'project_id': project_id,
            'trend': trend,
            'summary': {
                'cycles': len(trend),
                'latest_avg_final_score': PerformanceReviewService._round(float(trend[-1].get('avg_final_score') or 0), 2) if trend else 0,
                'latest_member_count': int(trend[-1].get('member_count') or 0) if trend else 0,
            }
        }

    @staticmethod
    def get_member_profile(project_id: int, member_name: str) -> Dict[str, Any]:
        trend_payload = PerformanceReviewService.get_member_trend(project_id, member_name, limit=12)
        latest = trend_payload['trend'][-1] if trend_payload['trend'] else {}

        with DatabasePool.get_connection() as conn:
            scorecard = conn.execute(DatabasePool.format_sql('''
                SELECT * FROM performance_score_cards
                WHERE project_id = ? AND target_member_name = ?
                ORDER BY generated_at DESC
                LIMIT 1
            '''), (project_id, member_name)).fetchone()
            scorecard = dict(scorecard) if scorecard else {}

            target = None
            if scorecard.get('cycle_id'):
                target = conn.execute(DatabasePool.format_sql('''
                    SELECT * FROM performance_review_targets
                    WHERE cycle_id = ? AND project_id = ? AND member_name = ?
                '''), (scorecard['cycle_id'], project_id, member_name)).fetchone()
            target = dict(target) if target else {}

            recognitions = conn.execute(DatabasePool.format_sql('''
                SELECT * FROM performance_recognition
                WHERE project_id = ? AND target_id = ?
                ORDER BY created_at DESC
                LIMIT 10
            '''), (project_id, target.get('id') or -1)).fetchall()
            recognitions = [dict(row) for row in recognitions]

            appeals = conn.execute(DatabasePool.format_sql('''
                SELECT * FROM performance_appeals
                WHERE scorecard_id = ?
                ORDER BY created_at DESC
                LIMIT 10
            '''), (scorecard.get('id') or -1,)).fetchall()
            appeals = [dict(row) for row in appeals]

            worklogs = conn.execute(DatabasePool.format_sql('''
                SELECT log_date, work_content, tomorrow_plan, issues_encountered, work_hours
                FROM work_logs
                WHERE project_id = ? AND member_name = ?
                ORDER BY log_date DESC
                LIMIT 8
            '''), (project_id, member_name)).fetchall()
            worklogs = [dict(row) for row in worklogs]

            member_row = conn.execute(DatabasePool.format_sql('''
                SELECT * FROM project_members
                WHERE project_id = ? AND name = ?
                LIMIT 1
            '''), (project_id, member_name)).fetchone()
            member_row = dict(member_row) if member_row else {}

        evidence = PerformanceReviewService._json_load(scorecard.get('evidence_json'), {})
        strengths = []
        if float(scorecard.get('warmth_score') or 0) >= 90:
            strengths.append('团队口碑稳定')
        if float(scorecard.get('onsite_score') or 0) >= 85:
            strengths.append('现场协作可靠')
        if float(scorecard.get('ai_score') or 0) >= 85:
            strengths.append('证据表现扎实')
        if evidence.get('delivery', {}).get('completed_tasks', 0) >= 3:
            strengths.append('交付兑现有连续性')
        if evidence.get('issues', {}).get('closed_owned_issues', 0) >= 2:
            strengths.append('问题闭环意识强')

        focus = []
        if evidence.get('issues', {}).get('high_pending_owned', 0) > 0:
            focus.append('优先处理名下高优问题')
        if float(scorecard.get('onsite_score') or 0) < 75:
            focus.append('补强现场沟通和协作透明度')
        if evidence.get('worklog', {}).get('completeness_ratio', 0) < 0.6:
            focus.append('提升日志与方案沉淀完整度')
        if not focus:
            focus.append('保持当前节奏，适合沉淀一条可复用经验')

        return {
            'project_id': project_id,
            'member_name': member_name,
            'member': member_row,
            'latest_scorecard': scorecard,
            'trend': trend_payload['trend'],
            'trend_summary': trend_payload['summary'],
            'strength_tags': strengths[:5],
            'focus_points': focus[:5],
            'recognitions': recognitions,
            'appeals': appeals,
            'recent_worklogs': worklogs,
            'latest_cycle': latest,
        }

    @staticmethod
    def update_cycle_status(cycle_id: int, status: str, current_user: Dict[str, Any]) -> Dict[str, Any]:
        role = str(current_user.get('role') or '')
        operator = str(current_user.get('display_name') or current_user.get('username') or '').strip() or '系统'
        if role not in ('admin', 'project_manager'):
            raise ValueError('当前账号没有周期锁定权限')
        status = str(status or '').strip()
        if status not in ('active', 'locked', 'archived'):
            raise ValueError('不支持的周期状态')

        with DatabasePool.get_connection() as conn:
            row = conn.execute(DatabasePool.format_sql('SELECT * FROM performance_review_cycles WHERE id = ?'), (cycle_id,)).fetchone()
            if not row:
                raise ValueError('未找到周评周期')
            cycle = dict(row)
            conn.execute(DatabasePool.format_sql('''
                UPDATE performance_review_cycles
                SET status = ?, updated_at = CURRENT_TIMESTAMP
                WHERE id = ?
            '''), (status, cycle_id))
            conn.commit()

        audit_service.log_operation(operator, '更新绩效周期状态', 'performance_cycle', cycle_id, cycle.get('title'), old_val={'status': cycle.get('status')}, new_val={'status': status})
        return {'success': True, 'status': status}

    @staticmethod
    def send_project_review_reminders(cycle_id: int, project_id: int, current_user: Dict[str, Any]) -> Dict[str, Any]:
        role = str(current_user.get('role') or '')
        operator = str(current_user.get('display_name') or current_user.get('username') or '').strip() or '系统'
        if role not in ('admin', 'project_manager'):
            raise ValueError('当前账号没有发送绩效提醒权限')

        with DatabasePool.get_connection() as conn:
            result = PerformanceReviewService._send_project_review_reminders_internal(conn, cycle_id, project_id, operator=operator)
            conn.commit()

        audit_service.log_operation(operator, '发送研发绩效评分提醒', 'performance_cycle', cycle_id, result.get('project_name'), new_val={'project_id': project_id, 'created': result.get('created', 0)})
        return result

    @staticmethod
    def export_cycle_markdown(cycle_id: int, current_user: Optional[Dict[str, Any]] = None, project_id: Optional[int] = None) -> Dict[str, Any]:
        data = PerformanceReviewService.get_overview(cycle_id=cycle_id, current_user=current_user or {}, project_id=project_id)
        cycle = data.get('cycle') or {}
        summary = data.get('summary') or {}
        lines = [
            f"# {cycle.get('title') or '研发绩效周评报告'}",
            "",
            f"- 周期：{cycle.get('start_date')} 至 {cycle.get('end_date')}",
            f"- 状态：{cycle.get('status') or '-'}",
            f"- 参评成员：{summary.get('member_count', 0)}",
            f"- 覆盖项目：{summary.get('project_count', 0)}",
            f"- 平均总分：{summary.get('avg_final_score', 0)}",
            f"- 待补评分：{summary.get('pending_reviews', 0)}",
            f"- 待处理申诉：{summary.get('pending_appeals', 0)}",
            "",
            "## 评分公式",
            "",
            "最终分 = 现场评价分×35% + AI证据分×55% + 人情味分×10% + 人工校准",
            "",
            "## 成员评分",
            "",
        ]

        for project in data.get('projects') or []:
            lines.append(f"### {project.get('project_name')} / {project.get('hospital_name') or '-'}")
            lines.append("")
            lines.append("| 成员 | 最终分 | 现场分 | AI分 | 人情味 | 评分数 | 状态 |")
            lines.append("| --- | ---: | ---: | ---: | ---: | ---: | --- |")
            for member in project.get('members') or []:
                scores = member.get('scores') or {}
                lines.append(
                    f"| {member.get('member_name')} | {scores.get('final', 0):.1f} | "
                    f"{scores.get('onsite', 0):.1f} | {scores.get('ai', 0):.1f} | "
                    f"{scores.get('warmth', 0):.1f} | {member.get('review_count', 0)} | "
                    f"{'已审批' if member.get('approved_by') else '未审批'} |"
                )
                if member.get('ai_highlight') or member.get('ai_support'):
                    lines.append(
                        f"| 评语 | {str(member.get('ai_highlight') or '').replace('|', ' ')} "
                        f"{str(member.get('ai_support') or '').replace('|', ' ')} |  |  |  |  |  |"
                    )
            lines.append("")

        if data.get('warm_wall'):
            lines.append("## 暖心记录")
            lines.append("")
            for item in data.get('warm_wall') or []:
                lines.append(f"- **{item.get('member_name')}**：{item.get('title') or ''} {item.get('content') or ''}")
            lines.append("")

        if data.get('pending_appeals'):
            lines.append("## 待处理申诉")
            lines.append("")
            for item in data.get('pending_appeals') or []:
                lines.append(f"- **{item.get('member_name')}**（{item.get('project_name')}）：{item.get('appeal_reason')}")
            lines.append("")

        scope_suffix = f"-project-{project_id}" if project_id else ''
        filename = f"performance-review-{cycle.get('cycle_key') or cycle_id}{scope_suffix}.md"
        return {
            'filename': filename,
            'content': "\n".join(lines),
        }

    @staticmethod
    def _ensure_temp_report_dir() -> str:
        temp_dir = os.path.abspath('temp_reports')
        os.makedirs(temp_dir, exist_ok=True)
        return temp_dir

    @staticmethod
    def _set_doc_normal_style(doc: Document):
        style = doc.styles['Normal']
        style.font.name = '微软雅黑'
        style.font.size = Pt(10.5)

    @staticmethod
    def export_cycle_docx(cycle_id: int, current_user: Optional[Dict[str, Any]] = None, project_id: Optional[int] = None) -> str:
        data = PerformanceReviewService.get_overview(cycle_id=cycle_id, current_user=current_user or {}, project_id=project_id)
        cycle = data.get('cycle') or {}
        summary = data.get('summary') or {}
        temp_dir = PerformanceReviewService._ensure_temp_report_dir()
        scope_suffix = f"-project-{project_id}" if project_id else ''
        filename = f"performance-review-{cycle.get('cycle_key') or cycle_id}{scope_suffix}.docx"
        file_path = os.path.join(temp_dir, filename)

        doc = Document()
        PerformanceReviewService._set_doc_normal_style(doc)

        title = doc.add_heading(cycle.get('title') or '研发绩效周评报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle = doc.add_paragraph(f"周期：{cycle.get('start_date')} 至 {cycle.get('end_date')} | 状态：{cycle.get('status') or '-'}")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_heading('一、总体概览', level=1)
        overview_table = doc.add_table(rows=4, cols=4)
        overview_table.style = 'Table Grid'
        overview_values = [
            ('参评成员', summary.get('member_count', 0), '覆盖项目', summary.get('project_count', 0)),
            ('平均总分', summary.get('avg_final_score', 0), '平均现场分', summary.get('avg_onsite_score', 0)),
            ('平均AI分', summary.get('avg_ai_score', 0), '待补评分', summary.get('pending_reviews', 0)),
            ('暖心记录', summary.get('recognition_count', 0), '待处理申诉', summary.get('pending_appeals', 0)),
        ]
        for row, values in zip(overview_table.rows, overview_values):
            for idx, value in enumerate(values):
                row.cells[idx].text = str(value)

        doc.add_heading('二、评分公式', level=1)
        doc.add_paragraph('最终分 = 现场评价分×35% + AI证据分×55% + 人情味分×10% + 人工校准')
        doc.add_paragraph('现场评价分 = 去极值后均值(责任心30% + 协作度25% + 响应速度25% + 专业度20%)')
        doc.add_paragraph('AI证据分 = 交付兑现40 + 问题闭环30 + 过程透明20 + 质量稳定10')
        doc.add_paragraph('人情味分 = 暖心卡、现场亮点、客户正反馈，只加不乱扣')

        doc.add_heading('三、成员评分明细', level=1)
        for project in data.get('projects') or []:
            doc.add_heading(f"{project.get('project_name')} / {project.get('hospital_name') or '-'}", level=2)
            table = doc.add_table(rows=1, cols=8)
            table.style = 'Table Grid'
            headers = ['成员', '角色', '最终分', '现场分', 'AI分', '人情味', '评分数', '状态']
            for idx, header in enumerate(headers):
                table.rows[0].cells[idx].text = header
            for member in project.get('members') or []:
                scores = member.get('scores') or {}
                row = table.add_row().cells
                row[0].text = str(member.get('member_name') or '')
                row[1].text = str(member.get('member_role') or '')
                row[2].text = f"{float(scores.get('final') or 0):.1f}"
                row[3].text = f"{float(scores.get('onsite') or 0):.1f}"
                row[4].text = f"{float(scores.get('ai') or 0):.1f}"
                row[5].text = f"{float(scores.get('warmth') or 0):.1f}"
                row[6].text = str(member.get('review_count') or 0)
                row[7].text = '已审批' if member.get('approved_by') else '未审批'
                if member.get('ai_highlight') or member.get('ai_support'):
                    doc.add_paragraph(f"{member.get('member_name')}：{member.get('ai_highlight') or ''} {member.get('ai_support') or ''}", style='List Bullet')

        doc.add_heading('四、暖心记录', level=1)
        if data.get('warm_wall'):
            for item in data.get('warm_wall') or []:
                doc.add_paragraph(f"{item.get('member_name')}：{item.get('title') or ''} {item.get('content') or ''}", style='List Bullet')
        else:
            doc.add_paragraph('本周期暂无暖心记录。')

        doc.add_heading('五、待处理申诉', level=1)
        if data.get('pending_appeals'):
            for item in data.get('pending_appeals') or []:
                doc.add_paragraph(f"{item.get('member_name')}（{item.get('project_name')}）：{item.get('appeal_reason')}", style='List Bullet')
        else:
            doc.add_paragraph('本周期暂无待处理申诉。')

        section = doc.sections[0]
        footer = section.footer.paragraphs[0]
        footer.text = f"此报告由 ICU-PM 绩效周评中心生成 | {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.save(file_path)
        return file_path

    @staticmethod
    def export_member_profile_docx(project_id: int, member_name: str) -> str:
        data = PerformanceReviewService.get_member_profile(project_id, member_name)
        temp_dir = PerformanceReviewService._ensure_temp_report_dir()
        safe_name = re.sub(r'[\\/:*?"<>|]+', '_', member_name or 'member')
        filename = f"performance-profile-{project_id}-{safe_name}.docx"
        file_path = os.path.join(temp_dir, filename)

        doc = Document()
        PerformanceReviewService._set_doc_normal_style(doc)

        title = doc.add_heading(f"{member_name} 绩效个人画像", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}").alignment = WD_ALIGN_PARAGRAPH.CENTER

        latest = data.get('latest_scorecard') or {}
        summary = data.get('trend_summary') or {}
        doc.add_heading('一、当前画像', level=1)
        table = doc.add_table(rows=2, cols=4)
        table.style = 'Table Grid'
        values = [
            ('当前最终分', f"{float(latest.get('final_score') or 0):.1f}", '历史平均', f"{float(summary.get('avg_final_score') or 0):.1f}"),
            ('历史最高', f"{float(summary.get('best_final_score') or 0):.1f}", '历史最低', f"{float(summary.get('lowest_final_score') or 0):.1f}"),
        ]
        for row, row_values in zip(table.rows, values):
            for idx, value in enumerate(row_values):
                row.cells[idx].text = str(value)

        doc.add_heading('二、优势标签', level=1)
        for item in data.get('strength_tags') or ['持续成长中']:
            doc.add_paragraph(str(item), style='List Bullet')

        doc.add_heading('三、当前关注点', level=1)
        for item in data.get('focus_points') or []:
            doc.add_paragraph(str(item), style='List Bullet')

        doc.add_heading('四、历史趋势', level=1)
        trend_table = doc.add_table(rows=1, cols=6)
        trend_table.style = 'Table Grid'
        for idx, header in enumerate(['周期', '现场分', 'AI分', '人情味分', '最终分', '状态']):
            trend_table.rows[0].cells[idx].text = header
        for item in data.get('trend') or []:
            row = trend_table.add_row().cells
            row[0].text = str(item.get('title') or '')
            row[1].text = f"{float(item.get('onsite_score') or 0):.1f}"
            row[2].text = f"{float(item.get('ai_score') or 0):.1f}"
            row[3].text = f"{float(item.get('warmth_score') or 0):.1f}"
            row[4].text = f"{float(item.get('final_score') or 0):.1f}"
            row[5].text = '已审批' if item.get('approved_by') else '未审批'

        doc.add_heading('五、近期暖心记录', level=1)
        if data.get('recognitions'):
            for item in data.get('recognitions') or []:
                doc.add_paragraph(f"{item.get('title') or '感谢你'}：{item.get('content') or ''}", style='List Bullet')
        else:
            doc.add_paragraph('暂无暖心记录。')

        doc.save(file_path)
        return file_path


performance_review_service = PerformanceReviewService()
