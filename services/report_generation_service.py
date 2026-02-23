import os
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from database import DatabasePool
from services.ai_service import AIService

class ReportGenerationService:
    @staticmethod
    def generate_formal_report(project_id, report_type="project_status"):
        """
        生成专业的项目正式 Word 报告
        :param project_id: 项目ID
        :param report_type: 报告类型 (project_status, phase_acceptance, risk_assessment)
        :return: 生成的文件路径
        """
        # 1. 获取项目全量数据
        data = ReportGenerationService._get_project_data(project_id)
        if not data:
            return None
        
        # 2. 调用 AI 生成专业摘要和分析
        ai_analysis = ReportGenerationService._get_ai_report_content(data, report_type)
        
        # 3. 创建 Word 文档
        temp_dir = 'temp_reports'
        if not os.path.exists(temp_dir):
            os.makedirs(temp_dir)
            
        file_name = f"{data['project']['project_name']}_{report_type}_{datetime.now().strftime('%Y%m%d%H%M')}.docx"
        file_path = os.path.abspath(os.path.join(temp_dir, file_name))
        
        doc = Document()
        
        # --- 文档样式设置 ---
        style = doc.styles['Normal']
        style.font.name = '微软雅黑'
        style.font.size = Pt(10.5)
        
        # --- 标题 ---
        title = doc.add_heading(data['project']['project_name'], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph(f"项目阶段性正式报告 ({datetime.now().strftime('%Y-%m-%d')})")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph("\n")
        
        # --- 基本信息表格 ---
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        
        cells = table.rows[0].cells
        cells[0].text = "项目名称"
        cells[1].text = data['project']['project_name']
        
        cells = table.rows[1].cells
        cells[0].text = "医院名称"
        cells[1].text = data['project']['hospital_name']
        
        cells = table.rows[2].cells
        cells[0].text = "项目经理"
        cells[1].text = data['project']['project_manager'] or "未指定"
        
        cells = table.rows[3].cells
        cells[0].text = "当前进度"
        cells[1].text = f"{data['project']['progress']}% ({data['project']['status']})"
        
        doc.add_paragraph("\n")
        
        # --- AI 智能分析部分 ---
        doc.add_heading('一、 AI 深度分析', level=1)
        doc.add_paragraph(ai_analysis)
        
        # --- 阶段任务详情 ---
        doc.add_heading('二、 项目阶段与任务明细', level=1)
        for stage in data['stages']:
            p = doc.add_paragraph()
            run = p.add_run(f"阶段：{stage['stage_name']} ({stage['progress']}%)")
            run.bold = True
            
            # 添加该阶段下的任务
            stage_tasks = [t for t in data['tasks'] if t['stage_id'] == stage['id']]
            if stage_tasks:
                for t in stage_tasks:
                    status = " [√]" if t['is_completed'] else " [ ]"
                    doc.add_paragraph(f"{status} {t['task_name']}", style='List Bullet')
            else:
                doc.add_paragraph("暂无明细任务", style='Normal')
        
        # --- 里程碑情况 ---
        doc.add_heading('三、 里程碑执行情况', level=1)
        if data['milestones']:
            m_table = doc.add_table(rows=1, cols=3)
            m_table.style = 'Table Grid'
            hdr_cells = m_table.rows[0].cells
            hdr_cells[0].text = '里程碑名称'
            hdr_cells[1].text = '计划日期'
            hdr_cells[2].text = '状态'
            
            for m in data['milestones']:
                row_cells = m_table.add_row().cells
                row_cells[0].text = m['name']
                row_cells[1].text = m['target_date'] or "-"
                row_cells[2].text = "已完成" if m['is_completed'] else "进行中"
        else:
            doc.add_paragraph("暂无关键里程碑记录")

        # --- 页脚 ---
        section = doc.sections[0]
        footer = section.footer
        p = footer.paragraphs[0]
        p.text = f"此报告由 ICU-PM AI 系统自动生成 | 生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.save(file_path)
        return file_path

    @staticmethod
    def _get_project_data(project_id):
        try:
            with DatabasePool.get_connection() as conn:
                project = conn.execute('SELECT * FROM projects WHERE id = ?', (project_id,)).fetchone()
                if not project:
                    return None
                
                stages = conn.execute('SELECT * FROM project_stages WHERE project_id = ? ORDER BY stage_order', (project_id,)).fetchall()
                milestones = conn.execute('SELECT * FROM milestones WHERE project_id = ? ORDER BY target_date', (project_id,)).fetchall()
                
                # 获取所有任务
                stage_ids = [s['id'] for s in stages]
                tasks = []
                if stage_ids:
                    placeholders = ', '.join(['?'] * len(stage_ids))
                    tasks = conn.execute(f'SELECT * FROM tasks WHERE stage_id IN ({placeholders})', stage_ids).fetchall()
                
                # 获取最近日志
                logs = conn.execute('SELECT * FROM work_logs WHERE project_id = ? ORDER BY log_date DESC LIMIT 20', (project_id,)).fetchall()
                
                return {
                    'project': dict(project),
                    'stages': [dict(s) for s in stages],
                    'milestones': [dict(m) for m in milestones],
                    'tasks': [dict(t) for t in tasks],
                    'logs': [dict(l) for l in logs]
                }
        except Exception as e:
            print(f"Get Project Data Error: {e}")
            return None

    @staticmethod
    def _get_ai_report_content(data, report_type):
        """调用 AI 生成报告摘要"""
        system_prompt = """你是一位资深的医疗信息化项目管理专家。
请根据提供的项目底层数据，生成一份结构化、专业、客观的“项目现状深度分析”。
要求：
1. 语言严谨，适合提交给医院领导或公司高层。
2. 包含：项目总体健康度评估、关键进展总结、潜在风险识别、下一步行动建议。
3. 请使用纯文本或 Markdown 格式（但不要包含代码块）。
4. 长度在 500-800 字左右。"""

        # 构建上下文
        context = f"项目名称: {data['project']['project_name']}\n"
        context += f"当前状态: {data['project']['status']}, 进度: {data['project']['progress']}%\n"
        context += f"里程碑详情: {[(m['name'], m['target_date'], m['is_completed']) for m in data['milestones']]}\n"
        
        recent_logs = "\n".join([f"- {l['log_date']}: {l['work_content']}" for l in data['logs']])
        context += f"最近工作进展:\n{recent_logs}\n"
        
        try:
            content = AIService.call_ai_api(system_prompt, context, task_type="analysis")
            return content or "AI 分析生成失败，请检查 AI 配置。"
        except Exception as e:
            return f"AI 接口调用异常: {str(e)}"

    @staticmethod
    def get_period_report_data(project_id, year, month=None, quarter=None, week=None):
        """
        获取特定时间段的项目报表数据
        """
        try:
            with DatabasePool.get_connection() as conn:
                project = conn.execute('SELECT * FROM projects WHERE id = ?', (project_id,)).fetchone()
                if not project:
                    return None
                
                # 构建时间筛选条件
                start_date = ""
                end_date = ""
                if month:
                    start_date = f"{year}-{int(month):02d}-01"
                    if int(month) == 12:
                        end_date = f"{int(year)+1}-01-01"
                    else:
                        end_date = f"{year}-{int(month)+1:02d}-01"
                elif quarter:
                    q_map = {1: ('01', '04'), 2: ('04', '07'), 3: ('07', '10'), 4: ('10', '01')}
                    start_m, end_m = q_map[int(quarter)]
                    start_date = f"{year}-{start_m}-01"
                    if int(quarter) == 4:
                        end_date = f"{int(year)+1}-01-01"
                    else:
                        end_date = f"{year}-{end_m}-01"
                elif week:
                    # 使用 ISO 周计算日期范围
                    d = datetime.fromisocalendar(int(year), int(week), 1)
                    start_date = d.strftime('%Y-%m-%d')
                    end_date = (d + timedelta(days=7)).strftime('%Y-%m-%d')
                
                # 获取该期间的日志
                logs = conn.execute(
                    'SELECT * FROM work_logs WHERE project_id = ? AND log_date >= ? AND log_date < ? ORDER BY log_date',
                    (project_id, start_date, end_date)
                ).fetchall()
                
                # 获取该期间完成的任务
                tasks = conn.execute(
                    '''SELECT t.*, s.stage_name FROM tasks t 
                       JOIN project_stages s ON t.stage_id = s.id 
                       WHERE s.project_id = ? AND t.is_completed = 1 AND t.completed_date >= ? AND t.completed_date < ?''',
                    (project_id, start_date, end_date)
                ).fetchall()
                
                # 获取该期间完成的里程碑
                milestones = conn.execute(
                    'SELECT * FROM milestones WHERE project_id = ? AND is_completed = 1 AND completed_date >= ? AND completed_date < ?',
                    (project_id, start_date, end_date)
                ).fetchall()
                
                # 获取财务数据 (汇总回款金额)
                revenue_res = conn.execute(
                    'SELECT SUM(amount) as collected_total FROM project_revenues WHERE project_id = ?', (project_id,)
                ).fetchone()
                collected_amount = revenue_res['collected_total'] if revenue_res['collected_total'] else 0
                
                return {
                    'project': dict(project),
                    'logs': [dict(l) for l in logs],
                    'tasks': [dict(t) for t in tasks],
                    'milestones': [dict(m) for m in milestones],
                    'financials': {
                        'contract_amount': project['contract_amount'] or 0,
                        'collected_amount': collected_amount
                    },
                    'period': {'year': year, 'month': month, 'quarter': quarter, 'week': week}
                }
        except Exception as e:
            print(f"Get Period Report Data Error: {e}")
            return None

    @staticmethod
    def generate_ai_business_summary(data):
        """调用 AI 生成业务层面的报表总结"""
        if not data: return "无数据供分析"
        
        system_prompt = """你是一位资深的医疗信息化项目群管理专家(PMO)。
请根据提供的项目月度/季度运行数据，撰写一份面向高层管理者的“业务运行总结”。
要求：
1. 语气专业、干练、结果导向。
2. 重点分析：本期核心达成（里程碑/任务）、执行效率评价、财务偏差分析、下期工作优先级建议。
3. 请使用 Markdown 格式排版。
4. 严禁罗列流水账，必须体现出分析和洞察。
5. 长度控制在 600 字左右。"""

        context = f"项目: {data['project']['project_name']} ({data['project']['hospital_name']})\n"
        period_str = f"{data['period']['year']}年"
        if data['period']['month']:
            period_str += f"{data['period']['month']}月"
        elif data['period']['quarter']:
            period_str += f"第{data['period']['quarter']}季度"
        elif data['period']['week']:
            period_str += f"第{data['period']['week']}周"
        
        context += f"分析周期: {period_str}\n"
        context += f"本期完成里程碑: {[m['name'] for m in data['milestones']]}\n"
        context += f"本期完成任务数: {len(data['tasks'])}\n"
        context += f"本期投入工时概况: {len(data['logs'])} 条工作内容记录\n"
        if data['financials']:
            context += f"财务快照: 合同总额 {data['financials']['contract_amount']}, 当前已回款 {data['financials']['collected_amount']}\n"
        
        try:
            content = AIService.call_ai_api(system_prompt, context, task_type="analysis")
            return content or "AI 分析生成失败。"
        except Exception as e:
            return f"AI 接口调用异常: {str(e)}"

report_gen_service = ReportGenerationService()
