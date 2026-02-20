"""从上传的 PDF/Word 文件中提取纯文本"""
import os

def extract_text_from_file(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == '.pdf':
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(file_path)
            return '\n'.join([page.extract_text() or '' for page in reader.pages])
        except Exception as e:
            return f'PDF解析失败: {e}'
    
    elif ext in ('.doc', '.docx'):
        try:
            from docx import Document
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs]
            # 也提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    paragraphs.append(' | '.join([cell.text for cell in row.cells]))
            return '\n'.join(paragraphs)
        except Exception as e:
            return f'Word解析失败: {e}'
    
    elif ext == '.txt':
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    return ''
